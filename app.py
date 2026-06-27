"""
Generador de informes de evaluación docente — EAFIT
Versión web (Streamlit)
"""

import io, os, re, zipfile, random, tempfile
import math
import json as _json
import urllib.request as _urlreq
import urllib.error as _urlerr
from collections import defaultdict
from copy import deepcopy
from lxml import etree
import openpyxl
import streamlit as st

# ─── CONFIGURACIÓN EXCEL ───────────────────────────────────────────────────
FILA_ENCABEZADO = 7
FILA_DATOS      = 8

COL_NOMBRE          = "Nombres y apellidos Docente"
COL_CICLO           = "Ciclo"
COL_CURSO           = "Nombre Catalogo"
COL_ID_ESCUELA      = "Id Escuela"
COL_ESCUELA         = "Escuela"
COL_COMPETENCIA     = "Competencia Evaluada"
COL_NOTA_FINAL      = "Nota competencia por clase"
COL_NOTA_FINAL_CLASE = "Nota final por clase"
COL_NOTA_CURSO      = "Nota final por curso"
COL_PREGUNTA        = "Pregunta"
COL_COMENTARIO      = "Comentarios"
COL_TOTAL_GENERADAS = "Total Evaluaciones generadas"
COL_EVALUACIONES    = "Evaluaciones realizadas"

ESCUELAS_EAFIT = {
    "E-ADM": "Escuela de Administración",
    "E-DER": "Escuela de Derecho",
    "E-ECO": "Escuela de Economía y Finanzas",
    "E-HUM": "Escuela de Humanidades",
    "E-ING": "Escuela de Ciencias Aplicadas e Ingeniería",
    "E-MED": "Escuela de Medicina",
    "E-MUS": "Escuela de Música",
    "E-DIS": "Escuela de Arquitectura y Diseño",
    "E-CS":  "Escuela de Ciencias",
    "E-VIS": "Vicerrectoría de Internacionalización",
}

FILTROS_COMENTARIOS = [
    # Respuestas vacías o sin contenido
    r"^\s*$",
    r"^\s*-+\s*$",
    r"^\s*\.+\s*$",
    # Afirmaciones/negaciones sin contenido
    r"^\s*(no|na|n\.a\.?|n/a|nada|ninguno?|ninguna?)\s*$",
    r"^\s*(nunguno?|nung[uo]no?)\s*$",
    r"^\s*(si|sí|yes)\s*$",
    r"^\s*(ok|oki|okay|okey)\s*$",
    # Frases cortas irrelevantes
    r"^\s*(todo\s+bien|todo\s+está?\s*bien|todo\s+esta\s*bien)\s*$",
    r"^\s*(bien|muy\s+bien|excelente|perfecto)\s*$",
    r"^\s*(gracias?|thanks?)\s*$",
    r"^\s*(no\s+aplica|no\s+apply|n\.?a\.?)\s*$",
    r"^\s*(cumple|cumplido|cumple\s+con\s+todo)\s*$",
    r"^\s*s[ií]n?\s+comentarios?\s*$",
    r"^\s*s[ií]n?\s+novedad(es)?\s*$",
    r"^\s*(ningún?\s+comentario|no\s+tengo\s+comentarios?)\s*$",
    r"^\s*(no\s+hay\s+comentarios?|sin\s+observaciones?)\s*$",
    r"^\s*(ningun[ao])\s*$",
]

MAYUSCULAS_FIJAS = {"eafit", "covid", "ia", "ti", "zoom", "teams", "meet", "canvas", "moodle"}

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# ─── (LanguageTool eliminado por rendimiento) ──────────────────────────────

# ─── HELPERS ───────────────────────────────────────────────────────────────

def resolver_escuela(id_escuela: str, escuela_raw: str) -> str:
    key = str(id_escuela or "").strip().upper()
    return ESCUELAS_EAFIT.get(key, str(escuela_raw or "").strip())

def fmt_nota(n) -> str:
    if n is None: return "—"
    try: return f"{float(n):.2f}".replace(".", ",")
    except: return str(n)

def es_valido(texto) -> bool:
    """Filtra comentarios sin contenido útil."""
    if not texto: return False
    t = str(texto).strip()
    for p in FILTROS_COMENTARIOS:
        if re.match(p, t, re.IGNORECASE): return False
    # Demasiado corto para ser útil
    if len(t) < 5:
        return False
    return True

def _sentence_case(texto: str) -> str:
    """
    Convierte el texto a sentence case inteligente:
    - Si todo (o casi todo) está en mayúsculas → convierte a minúsculas primero
    - Primera letra de la oración en mayúscula
    - Respeta siglas conocidas (EAFIT, COVID, etc.)
    - Respeta palabras con mayúscula interna (iPad, etc.) — no las toca
    """
    if not texto:
        return texto

    # Detectar si el texto está "gritado" (>55% letras en mayúscula)
    letras = [c for c in texto if c.isalpha()]
    if letras and sum(1 for c in letras if c.isupper()) / len(letras) > 0.55:
        texto = texto.lower()

    # Procesar palabra por palabra
    palabras = texto.split()
    resultado = []
    for i, palabra in enumerate(palabras):
        base = palabra.rstrip(".,;:!?()")
        sufijo = palabra[len(base):]
        base_lower = base.lower()

        if base_lower in MAYUSCULAS_FIJAS:
            resultado.append(base.upper() + sufijo)
        elif len(base) > 1 and base.isupper():
            # Sigla desconocida → dejar en mayúsculas
            resultado.append(palabra)
        else:
            resultado.append(base_lower + sufijo)

    if resultado:
        p = resultado[0]
        resultado[0] = p[0].upper() + p[1:] if p else p

    texto_f = " ".join(resultado)

    # Asegurar punto al final
    if texto_f and texto_f[-1] not in ".!?;:":
        texto_f += "."

    return texto_f

def formatear_comentario(texto: str) -> str:
    """Sentence case únicamente (LanguageTool eliminado por rendimiento)."""
    texto = str(texto).strip()
    if not texto:
        return texto
    texto = _sentence_case(texto)
    if texto:
        texto = texto[0].upper() + texto[1:]
    return texto

def slugify(texto: str) -> str:
    repl = {'á':'a','é':'e','í':'i','ó':'o','ú':'u','ü':'u','ñ':'n',
            'Á':'A','É':'E','Í':'I','Ó':'O','Ú':'U','Ü':'U','Ñ':'N'}
    for orig, rep in repl.items():
        texto = texto.replace(orig, rep)
    texto = re.sub(r'[^\w\s\-]', '', texto)
    return re.sub(r'\s+', '', texto)

def nombre_archivo_defecto(datos: dict, nombre_prof: str) -> str:
    info       = datos.get("info", {})
    ciclo      = str(info.get("ciclo", "")).strip()
    escuela_id = str(info.get("id_escuela", "")).replace("-", "")
    curso      = slugify(info.get("curso", ""))
    partes     = nombre_prof.strip().split()
    if len(partes) >= 2:
        mitad          = len(partes) // 2
        primer_nombre  = partes[mitad].capitalize()
        primer_apellido= partes[0].capitalize()
        nombre_corto   = primer_nombre + primer_apellido
    else:
        nombre_corto = slugify(nombre_prof)
    return f"{ciclo}_{escuela_id}_{curso}_{nombre_corto}"

def replace_in_subtree(elem, old: str, new: str):
    for t in elem.iter(W+'t'):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)

def clone_bullet_para(template_para):
    new_p = deepcopy(template_para)
    for attr in list(new_p.attrib):
        if 'paraId' in attr or 'textId' in attr:
            new_p.set(attr, f"{random.randint(0x10000000, 0xFFFFFFFF):08X}")
    return new_p

# ─── LECTURA EXCEL ─────────────────────────────────────────────────────────

COL_CATALOGO = "Catálogo"
COL_NCLASE   = "Nº Clase"

def leer_excel(archivo_bytes: bytes = None,
               filtro_catalogo: str = None,
               filtro_clase: str = None,
               archivo_path: str = None) -> dict:
    """
    Lee el Excel de evaluaciones docentes y filtra por catálogo/clase.
    Usa _cargar_filas_excel (cacheada) para no releer todo el archivo en cada búsqueda.
    """
    filas_por_clave, headers = _cargar_filas_excel(archivo_path, archivo_bytes)

    f_cat   = str(filtro_catalogo).strip().upper() if filtro_catalogo else None
    f_clase = str(filtro_clase).strip() if filtro_clase else None
    clave   = (f_cat or "", f_clase or "")

    filas = filas_por_clave.get(clave, [])
    return _procesar_filas(filas, headers)


@st.cache_data(show_spinner=False)
def _cargar_filas_excel(archivo_path: str = None, archivo_bytes: bytes = None):
    """
    Lee el Excel UNA sola vez y agrupa las filas por (catálogo, nº clase).
    Resultado cacheado: las búsquedas siguientes son instantáneas (no se vuelve a leer el Excel).
    El caché se invalida automáticamente si cambia el archivo (por su contenido/ruta+mtime).
    """
    if archivo_path:
        wb = openpyxl.load_workbook(archivo_path, read_only=True)
    else:
        wb = openpyxl.load_workbook(io.BytesIO(archivo_bytes), read_only=True)
    ws = wb.active

    headers = {}
    for row in ws.iter_rows(min_row=FILA_ENCABEZADO, max_row=FILA_ENCABEZADO, values_only=True):
        for i, val in enumerate(row):
            if val: headers[str(val).strip()] = i
        break

    idx_cat   = headers.get(COL_CATALOGO)
    idx_clase = headers.get(COL_NCLASE)
    idx_nombre = headers.get(COL_NOMBRE)

    filas_por_clave = defaultdict(list)
    for row in ws.iter_rows(min_row=FILA_DATOS, values_only=True):
        if idx_nombre is None or idx_nombre >= len(row) or not row[idx_nombre]:
            continue
        cat_row   = str(row[idx_cat] if idx_cat is not None and idx_cat < len(row) else "").strip().upper()
        clase_row = str(row[idx_clase] if idx_clase is not None and idx_clase < len(row) else "").strip()
        filas_por_clave[(cat_row, clase_row)].append(row)

    wb.close()
    return dict(filas_por_clave), headers


def _procesar_filas(rows, headers) -> dict:
    """Construye el diccionario de profesores a partir de las filas ya filtradas."""
    def col(row_data, nombre):
        idx = headers.get(nombre)
        return row_data[idx] if idx is not None and idx < len(row_data) else None

    profesores = defaultdict(lambda: {
        "info": {}, "nota_final": None, "nota_curso": None,
        "comentarios": defaultdict(list),
        "total_generadas": None, "evaluaciones_realizadas": None,
        "notas_competencias": {},
    })

    for row in rows:
        nombre = col(row, COL_NOMBRE)
        if not nombre: continue
        nombre = str(nombre).strip()
        p = profesores[nombre]

        if not p["info"]:
            id_esc  = str(col(row, COL_ID_ESCUELA) or "").strip()
            esc_raw = str(col(row, COL_ESCUELA) or "").strip()
            cat     = str(col(row, COL_CATALOGO) or "").strip().lstrip()
            nclase  = str(col(row, COL_NCLASE) or "").strip()
            modo    = str(col(row, "Modo de Enseñanza") or "").strip()
            p["info"] = {
                "ciclo":           str(col(row, COL_CICLO) or "").strip(),
                "curso":           str(col(row, COL_CURSO) or "").strip(),
                "id_escuela":      id_esc,
                "escuela":         resolver_escuela(id_esc, esc_raw),
                "catalogo_clase":  f"{cat}-{nclase}" if cat and nclase else "",
                "modalidad":       modo,
            }

        comp      = str(col(row, COL_COMPETENCIA) or "").strip()
        pregunta  = str(col(row, COL_PREGUNTA) or "").strip()
        nota_final= col(row, COL_NOTA_FINAL)
        nota_final_clase = col(row, COL_NOTA_FINAL_CLASE)
        nota_curso= col(row, COL_NOTA_CURSO)
        comentario= col(row, COL_COMENTARIO)

        total_gen = col(row, COL_TOTAL_GENERADAS)
        eval_real = col(row, COL_EVALUACIONES)
        if total_gen and p["total_generadas"] is None:
            try: p["total_generadas"] = int(float(total_gen))
            except: pass
        if eval_real and p["evaluaciones_realizadas"] is None:
            try: p["evaluaciones_realizadas"] = int(float(eval_real))
            except: pass

        if nota_final and nota_final > 0 and p["nota_final"] is None:
            p["nota_final"] = nota_final
        if nota_curso and nota_curso > 0 and p["nota_curso"] is None:
            p["nota_curso"] = nota_curso

        # Puntaje de la competencia para el diagrama de araña:
        # algunas competencias (Índice de recomendación, Pacto Pedagógico, Calidad,
        # Comportamiento del grupo) vienen en 0 en "Nota competencia por clase" pero
        # sí traen el valor real en "Nota final por clase" → se usa como fallback.
        nota_para_grafico = nota_final if (nota_final and nota_final > 0) else nota_final_clase
        if comp and comp != "Comentarios" and nota_para_grafico and nota_para_grafico > 0:
            if comp not in p["notas_competencias"]:
                p["notas_competencias"][comp] = nota_para_grafico

        if comp == "Comentarios" and pregunta and comentario:
            if es_valido(str(comentario)):
                p["comentarios"][pregunta].append(formatear_comentario(str(comentario)))

    return dict(profesores)

# ─── GENERACIÓN WORD EN MEMORIA ────────────────────────────────────────────

def generar_informe_bytes(nombre: str, datos: dict,
                          plantilla_bytes: bytes,
                          nombre_archivo: str = None) -> tuple[bytes, str]:
    """Devuelve (docx_bytes, nombre_archivo). Compatible con Plantilla2 (placeholders explícitos)."""
    info                    = datos["info"]
    nota_curso              = datos["nota_curso"]
    nota_final              = datos["nota_final"]
    comentarios             = datos["comentarios"]
    total_generadas         = datos.get("total_generadas")
    evaluaciones_realizadas = datos.get("evaluaciones_realizadas")

    if not nombre_archivo:
        nombre_archivo = nombre_archivo_defecto(datos, nombre)
    nombre_archivo = re.sub(r'[<>:"/\\|?*]', '_', nombre_archivo).strip()

    tree = etree.fromstring(
        zipfile.ZipFile(io.BytesIO(plantilla_bytes)).read('word/document.xml')
    )
    body = tree.find(W+'body')

    def texto_de(elem):
        return ''.join(t.text or '' for t in elem.iter(W+'t')).strip()

    def reemplazar(elem, placeholder: str, valor: str):
        """
        Reemplaza placeholder en el texto concatenado de todos los runs de un elemento.
        Word frecuentemente fragmenta el texto en múltiples <w:r> y <w:t>, por lo que
        buscar placeholder en cada <w:t> individual falla. Esta función consolida el texto
        de cada párrafo en un único run antes de reemplazar, preservando el estilo del
        primer run original.
        """
        for para in elem.iter(W+'p'):
            runs = para.findall('.//' + W+'r')
            if not runs:
                continue
            # Texto completo del párrafo concatenando todos los <w:t>
            texto_completo = ''.join(t.text or '' for r in runs for t in r.findall(W+'t'))
            if placeholder not in texto_completo:
                continue
            # Reemplazar el placeholder
            nuevo_texto = texto_completo.replace(placeholder, valor)
            # Tomar el estilo (rPr) del primer run
            primer_run = runs[0]
            rpr = primer_run.find(W+'rPr')
            # Eliminar todos los runs del párrafo
            for r in runs:
                para.remove(r)
            # Crear un único run nuevo con el texto reemplazado
            new_r = etree.SubElement(para, W+'r')
            if rpr is not None:
                new_r.insert(0, deepcopy(rpr))
            new_t = etree.SubElement(new_r, W+'t')
            new_t.text = nuevo_texto
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # ── Tasa de respuesta ──
    tasa_str = ""
    if total_generadas and evaluaciones_realizadas and total_generadas > 0:
        tasa_str = f"{round(evaluaciones_realizadas / total_generadas * 100)}%"

    # ── Portada: placeholders directos ──
    reemplazar(body, '{Nombre_del_curso}', info.get('curso', ''))
    reemplazar(body, '{Nombre_programa}',  info.get('escuela', ''))
    reemplazar(body, '{Numero_codigo}',    info.get('catalogo_clase', ''))
    reemplazar(body, '{Numero_semestre}',  info.get('ciclo', ''))
    reemplazar(body, '{Nombre_profesor}',  nombre.title())

    # ── Tasa de respuesta: párrafo "Tasa de respuesta: %" → "Tasa de respuesta: 47%" ──
    for child in list(body):
        if 'Tasa de respuesta' in texto_de(child):
            reemplazar(child, '%', tasa_str)
            break

    # ── Tabla estudiantes: placeholders {{number_students_answered}} / {{total_number_students}} ──
    reemplazar(body, '{{number_students_answered}}',
               str(evaluaciones_realizadas) if evaluaciones_realizadas is not None else '—')
    reemplazar(body, '{{total_number_students}}',
               str(total_generadas) if total_generadas is not None else '—')

    # ── Diagrama de araña: párrafo {{Spider_diagram}} → imagen PNG ──
    notas_comp_raw = datos.get("notas_competencias", {})
    ORDEN_COMPETENCIAS = [
        "Relacional (Calidad Humana)",
        "Pedagógica",
        "Relacional (Respeto)",
        "Índice de recomendación",
        "Pacto Pedagógico",
    ]
    notas_comp = {k: notas_comp_raw[k] for k in ORDEN_COMPETENCIAS if k in notas_comp_raw}

    _spider_png_data = _spider_img_rid = _spider_img_target = None

    p_spider = next((c for c in list(body) if '{{Spider_diagram}}' in texto_de(c)), None)
    if p_spider is not None and notas_comp and len(notas_comp) >= 3:
        png_bytes  = _spider_chart_png(notas_comp)
        img_rid    = "rIdSpider"
        img_name   = "spider_chart.png"
        img_target = f"media/{img_name}"
        EMU_W = EMU_H = 5000000   # ~13.9 cm cuadrado

        draw_xml = (
            f'<w:p xmlns:w="{W[1:-1]}"'
            f' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
            f' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
            f' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
            f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<w:pPr><w:jc w:val="center"/></w:pPr>'
            f'<w:r><w:rPr/><w:drawing>'
            f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
            f'<wp:extent cx="{EMU_W}" cy="{EMU_H}"/>'
            f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            f'<wp:docPr id="1" name="SpiderChart"/>'
            f'<wp:cNvGraphicFramePr>'
            f'<a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>'
            f'</wp:cNvGraphicFramePr>'
            f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
            f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            f'<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            f'<pic:nvPicPr><pic:cNvPr id="0" name="{img_name}"/><pic:cNvPicPr/></pic:nvPicPr>'
            f'<pic:blipFill><a:blip r:embed="{img_rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
            f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{EMU_W}" cy="{EMU_H}"/></a:xfrm>'
            f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
            f'</pic:pic></a:graphicData></a:graphic>'
            f'</wp:inline></w:drawing></w:r></w:p>'
        )
        img_para = etree.fromstring(draw_xml)
        pos = list(body).index(p_spider)
        body.remove(p_spider)
        body.insert(pos, img_para)
        _spider_png_data   = png_bytes
        _spider_img_rid    = img_rid
        _spider_img_target = img_target
    elif p_spider is not None:
        reemplazar(p_spider, '{{Spider_diagram}}', '')

    # ── Tabla competencias: eliminar (reemplazada por el diagrama de araña) ──
    for child in list(body):
        if child.tag != W+'tbl':
            continue
        rows_tbl = child.findall('.//' + W+'tr')
        if not rows_tbl:
            continue
        header = texto_de(rows_tbl[0])
        if 'competencias' not in header.lower():
            continue
        body.remove(child)
        break

    # ── Comentarios: reemplazar {Comentario} (con viñeta heredada del template) ──
    def get_comentarios(clave):
        for k, lista in comentarios.items():
            if clave in k.lower():
                return lista
        return []

    for titulo_frag, clave in [("positivo","positivo"),("mejorar","mejorar"),("adicional","adicional")]:
        body_children = list(body)
        idx_titulo = next((i for i, c in enumerate(body_children)
                           if titulo_frag in texto_de(c).lower()), None)
        if idx_titulo is None:
            continue
        placeholders = []
        for c in body_children[idx_titulo + 1:]:
            t = texto_de(c)
            if t == '{Comentario}':
                placeholders.append(c)
            elif t and t != '{Comentario}':
                break
        if not placeholders:
            continue
        template_para = placeholders[0]
        insert_pos    = list(body).index(placeholders[0])
        for pp in placeholders:
            body.remove(pp)
        for j, texto in enumerate(get_comentarios(clave)):
            new_p  = clone_bullet_para(template_para)
            all_ts = list(new_p.iter(W+'t'))
            if all_ts:
                all_ts[0].text = texto
                all_ts[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                for t in all_ts[1:]:
                    t.text = ''
            body.insert(insert_pos + j, new_p)

    # ── Consideraciones: limpiar placeholders de ambas versiones de plantilla ──
    reemplazar(body, '{Comentario parrafo}', '')   # plantilla vieja
    for child in list(body):
        t = texto_de(child)
        if '{ Comentario' in t and 'Profesor' in t:
            body.remove(child)                      # plantilla nueva

    # ── POST-PROCESO: ajustes de estilo ──────────────────────────────────────

    # Cambio 3: Tabla estudiantes — igualar fuente sz=18 y alineación center en ambas filas
    for child in list(body):
        if child.tag != W+'tbl':
            continue
        if 'estudiantes' not in texto_de(child).lower():
            continue
        for row in child.findall('.//' + W+'tr'):
            cells = row.findall(W+'tc')
            if len(cells) < 2:
                continue
            cell_valor = cells[1]
            # Asegurar párrafo con alineación centrada
            for p in cell_valor.findall('.//' + W+'p'):
                pPr = p.find(W+'pPr')
                if pPr is None:
                    pPr = etree.SubElement(p, W+'pPr')
                    p.insert(0, pPr)
                jc = pPr.find(W+'jc')
                if jc is None:
                    jc = etree.SubElement(pPr, W+'jc')
                jc.set(W+'val', 'center')
            # Asegurar sz=18 en todos los runs de la celda valor
            for r in cell_valor.findall('.//' + W+'r'):
                rPr = r.find(W+'rPr')
                if rPr is None:
                    rPr = etree.Element(W+'rPr')
                    r.insert(0, rPr)
                for tag in [W+'sz', W+'szCs']:
                    el = rPr.find(tag)
                    if el is None:
                        el = etree.SubElement(rPr, tag)
                    el.set(W+'val', '18')
        break

    # Cambio 4: Tabla competencias — ampliar para que "Relacional (Calidad Humana)" quede en un renglón
    for child in list(body):
        if child.tag != W+'tbl':
            continue
        if 'competencias' not in texto_de(child).lower():
            continue
        # Ampliar ancho total de la tabla
        tblPr = child.find(W+'tblPr')
        if tblPr is not None:
            tblW = tblPr.find(W+'tblW')
            if tblW is None:
                tblW = etree.SubElement(tblPr, W+'tblW')
            tblW.set(W+'w', '4600')
            tblW.set(W+'type', 'dxa')
        # Ajustar anchos de cada celda
        for row in child.findall('.//' + W+'tr'):
            cells = row.findall(W+'tc')
            if len(cells) < 2:
                continue
            for ci, (cell, new_w) in enumerate(zip(cells, ['3800', '800'])):
                tcPr = cell.find(W+'tcPr')
                if tcPr is None:
                    tcPr = etree.SubElement(cell, W+'tcPr')
                    cell.insert(0, tcPr)
                tcW = tcPr.find(W+'tcW')
                if tcW is None:
                    tcW = etree.SubElement(tcPr, W+'tcW')
                tcW.set(W+'w', new_w)
                tcW.set(W+'type', 'dxa')
        break

    # Cambio 5: Estilos de sección de comentarios
    # "Preguntas abiertas" → azul EAFIT (#0B4DFF), bold
    # Títulos de cada sección (Menciona positivo / Menciona mejorar / ¿Tienes algún...) → bold
    COLOR_AZUL = '0B4DFF'
    for child in list(body):
        t = texto_de(child)
        if t == 'Preguntas abiertas':
            for r in child.findall('.//' + W+'r'):
                rPr = r.find(W+'rPr')
                if rPr is None:
                    rPr = etree.Element(W+'rPr')
                    r.insert(0, rPr)
                # Bold
                if rPr.find(W+'b') is None:
                    etree.SubElement(rPr, W+'b')
                # Color azul
                color_el = rPr.find(W+'color')
                if color_el is None:
                    color_el = etree.SubElement(rPr, W+'color')
                color_el.set(W+'val', COLOR_AZUL)
        elif any(frag in t.lower() for frag in ['menciona un aspecto', '¿tienes algún comentario']):
            for r in child.findall('.//' + W+'r'):
                rPr = r.find(W+'rPr')
                if rPr is None:
                    rPr = etree.Element(W+'rPr')
                    r.insert(0, rPr)
                if rPr.find(W+'b') is None:
                    etree.SubElement(rPr, W+'b')

    # Cambio 6: Eliminar completamente el párrafo de Modalidad
    for child in list(body):
        if 'Modalidad:' in texto_de(child):
            body.remove(child)
            break

    # ── Empaquetar ──

    new_xml = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
    out_buf  = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(plantilla_bytes), 'r') as zin:
        with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/document.xml':
                    zout.writestr(item, new_xml)
                elif _spider_png_data and item.filename == 'word/_rels/document.xml.rels':
                    # Inyectar la relationship de la imagen
                    rels_xml = zin.read(item.filename)
                    rels_tree = etree.fromstring(rels_xml)
                    REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
                    new_rel = etree.SubElement(rels_tree, f"{{{REL_NS}}}Relationship")
                    new_rel.set("Id", _spider_img_rid)
                    new_rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")
                    new_rel.set("Target", _spider_img_target)
                    zout.writestr(item, etree.tostring(rels_tree, xml_declaration=True,
                                                       encoding='UTF-8', standalone=True))
                else:
                    zout.writestr(item, zin.read(item.filename))
            # Escribir el PNG de la araña dentro del zip
            if _spider_png_data:
                zout.writestr(f"word/{_spider_img_target}", _spider_png_data)
    return out_buf.getvalue(), nombre_archivo

# ─── SPIDER CHART PNG (para insertar en Word) ───────────────────────────────

def _spider_chart_png(notas: dict) -> bytes:
    """
    Genera el diagrama de araña como PNG en memoria usando matplotlib.
    Incluye todas las competencias + índice de recomendación + pacto pedagógico si existen.
    Retorna los bytes del PNG.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np
    import textwrap

    # ── Colores EAFIT ──────────────────────────────────────────────────────
    C_BLUE      = "#004B85"
    C_BLUE_MID  = "#0064B0"
    C_YELLOW    = "#FFB903"
    C_FILL      = "#004B85"
    C_GRID      = "#DDDDDD"
    C_BG        = "#FFFFFF"
    C_LABEL     = "#1A1A2E"
    C_SCORE_BG  = "#004B85"
    C_SCORE_FG  = "#FFFFFF"

    labels_raw = list(notas.keys())
    values     = [min(max(float(v), 0), 5) for v in notas.values()]
    n          = len(labels_raw)

    # Abreviar etiquetas largas para que quepan en el radar
    def _abrev(txt: str, max_chars: int = 22) -> str:
        txt = txt.strip()
        if len(txt) <= max_chars:
            return txt
        # Dividir en hasta 2 líneas
        words = txt.split()
        lines, current = [], ""
        for w in words:
            if len(current) + len(w) + 1 <= max_chars:
                current = (current + " " + w).strip()
            else:
                if current:
                    lines.append(current)
                current = w
        if current:
            lines.append(current)
        return "\n".join(lines[:2])

    labels = [_abrev(l) for l in labels_raw]

    # ── Geometría ──────────────────────────────────────────────────────────
    angles      = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    angles     += angles[:1]
    values_plot = values + values[:1]

    # Canvas más grande para que las etiquetas tengan espacio
    fig = plt.figure(figsize=(8, 8), facecolor=C_BG)
    ax  = fig.add_subplot(111, polar=True, facecolor="#F8FAFD")

    # ── Rejilla ────────────────────────────────────────────────────────────
    ax.set_ylim(0, 5)
    levels = [1, 2, 3, 4, 5]
    ax.set_yticks(levels)
    ax.set_yticklabels([])           # ocultamos los ticks de radio
    ax.yaxis.set_visible(False)

    # Anillos concéntricos dibujados a mano (más control visual)
    for lvl in levels:
        ring_angles = np.linspace(0, 2 * np.pi, 300)
        ring_r      = [lvl] * 300
        alpha       = 0.55 if lvl == 5 else 0.25
        lw          = 1.2  if lvl == 5 else 0.6
        ax.plot(ring_angles, ring_r, color=C_GRID, linewidth=lw,
                linestyle="-", alpha=alpha, zorder=1)
        # Número del nivel en el eje 0 (arriba)
        if lvl < 5:
            ax.text(0, lvl + 0.07, str(lvl), ha="center", va="bottom",
                    fontsize=7, color="#AAAAAA", zorder=5)

    # Líneas de los ejes (radios)
    for angle in angles[:-1]:
        ax.plot([angle, angle], [0, 5], color=C_GRID,
                linewidth=0.8, alpha=0.5, zorder=1)

    # Ocultar la rejilla y bordes por defecto de matplotlib
    ax.grid(False)
    ax.spines["polar"].set_visible(False)

    # ── Polígono de datos ──────────────────────────────────────────────────
    ax.fill(angles, values_plot, color=C_FILL, alpha=0.18, zorder=2)
    ax.plot(angles, values_plot, color=C_BLUE, linewidth=2.5,
            linestyle="solid", zorder=3)

    # ── Puntos en cada vértice ─────────────────────────────────────────────
    for angle, val in zip(angles[:-1], values):
        ax.plot(angle, val, "o", color=C_YELLOW, markersize=10,
                markeredgecolor=C_BLUE, markeredgewidth=1.8, zorder=4)

    # ── Etiquetas de puntaje con pastilla de fondo ─────────────────────────
    # Desplazamiento adaptativo: alejar del centro
    for angle, val, lbl_raw in zip(angles[:-1], values, labels_raw):
        offset = 0.52   # distancia adicional sobre el punto
        r_text = min(val + offset, 5.0)

        # Evitar que quede fuera del área si el valor es muy alto
        if val > 4.4:
            r_text = val - 0.45

        score_str = f"{val:.2f}"
        ax.annotate(
            score_str,
            xy=(angle, val),
            xytext=(angle, r_text),
            textcoords="data",
            ha="center", va="center",
            fontsize=8.5, fontweight="bold", color=C_SCORE_FG,
            bbox=dict(
                boxstyle="round,pad=0.25",
                facecolor=C_SCORE_BG,
                edgecolor=C_YELLOW,
                linewidth=1.2,
                alpha=0.92,
            ),
            zorder=6,
        )

    # ── Etiquetas de los ejes (categorías) ────────────────────────────────
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels([])   # usamos anotaciones manuales para control total

    label_r = 5.65   # radio base donde ponemos las etiquetas
    for i, (angle, label) in enumerate(zip(angles[:-1], labels)):
        ha = "center"
        if 0.1 < angle < np.pi - 0.1:
            ha = "left" if angle < np.pi else "right"
        if np.pi + 0.1 < angle < 2 * np.pi - 0.1:
            ha = "right"

        # Calcular x, y en coordenadas cartesianas para anotación fuera del polar
        x = np.sin(angle)   # matplotlib polar: 0° arriba, gira horario
        y = np.cos(angle)

        ax.text(
            angle, label_r, label,
            ha="center", va="center",
            fontsize=8.2, color=C_LABEL, fontweight="600",
            multialignment="center",
            linespacing=1.3,
            zorder=7,
        )

    # ── Tick params ───────────────────────────────────────────────────────
    ax.tick_params(axis="both", which="both",
                   bottom=False, top=False, left=False, right=False,
                   labelbottom=False, labeltop=False)

    # ── Ajuste márgenes para que las etiquetas no se corten ───────────────
    ax.set_rlim(0, 5)
    fig.subplots_adjust(left=0.08, right=0.92, top=0.88, bottom=0.08)

    # ── Título ────────────────────────────────────────────────────────────
    fig.text(
        0.5, 0.96,
        "Resultados por competencia",
        ha="center", va="top",
        fontsize=11, fontweight="bold", color=C_BLUE,
    )

    # ── Leyenda de puntaje máximo ─────────────────────────────────────────
    patch = mpatches.Patch(facecolor=C_FILL, alpha=0.35,
                           edgecolor=C_BLUE, linewidth=1.2,
                           label="Puntaje obtenido (escala 0–5)")
    ax.legend(handles=[patch], loc="lower center",
              bbox_to_anchor=(0.5, -0.08), frameon=False,
              fontsize=8, labelcolor="#555555")

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor=C_BG, edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ─── MÓDULO: ALISTAMIENTO DE CONSIDERACIONES (IA) ─────────────────────────
#
# Este módulo toma un informe .docx YA GENERADO por la app y, usando uno o
# varios documentos guía (formaciones EXA, protocolos, lineamientos, etc.)
# como contexto, reescribe ÚNICAMENTE la sección "Consideraciones" del
# informe usando GitHub Models (API gratuita, modelo gpt-4o-mini).
#
# El resto del informe (portada, diagrama de competencias, comentarios,
# aspectos formativos, firma) no se toca en absoluto.

GITHUB_MODELS_URL   = "https://models.github.ai/inference/chat/completions"
GITHUB_MODELS_MODEL = "openai/gpt-4o"
MAX_CHARS_CONTEXTO  = 16000   # ~4000 tokens — margen seguro para gpt-4o (límite: 8000 tokens)

@st.cache_data(show_spinner=False)
def extraer_texto_referencia(nombre_archivo: str, contenido: bytes) -> str:
    """Extrae texto plano de un archivo de referencia (.pdf, .docx, .txt). Resultado cacheado."""
    ext = nombre_archivo.lower().rsplit(".", 1)[-1] if "." in nombre_archivo else ""
    try:
        if ext == "txt":
            return contenido.decode("utf-8", errors="ignore")

        elif ext == "docx":
            import docx as _docx_lib
            doc = _docx_lib.Document(io.BytesIO(contenido))
            partes = [p.text for p in doc.paragraphs if p.text.strip()]
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        if celda.text.strip():
                            partes.append(celda.text.strip())
            return "\n".join(partes)

        elif ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(contenido))
            partes = []
            for pagina in reader.pages:
                texto = pagina.extract_text() or ""
                if texto.strip():
                    partes.append(texto)
            return "\n".join(partes)

        else:
            return ""
    except Exception:
        return ""


def extraer_texto_informe_actual(docx_bytes: bytes) -> dict:
    """
    Extrae del informe .docx: portada, comentarios, puntajes de competencias
    y estado de la sección Consideraciones.
    """
    tree = etree.fromstring(
        zipfile.ZipFile(io.BytesIO(docx_bytes)).read('word/document.xml')
    )
    body = tree.find(W + 'body')
    children = list(body)

    def texto_de(elem):
        return "".join(t.text or "" for t in elem.iter(W + 't')).strip()

    textos = [texto_de(c) for c in children]

    # Portada: buscar párrafos que empiecen con los labels conocidos
    labels_portada = ["Nombre del curso", "Nombre del programa", "Código curso",
                      "Semestre", "Modalidad", "Nombre profesor", "Tasa de respuesta"]
    portada_lineas = []
    for t in textos:
        if any(t.startswith(lbl) for lbl in labels_portada):
            portada_lineas.append(t)
    portada = "\n".join(portada_lineas) if portada_lineas else (textos[0] if textos else "")

    # Competencias: buscar en tabla (texto con "Competencias evaluadas")
    competencias_lineas = []
    for child in children:
        if child.tag == W + 'tbl':
            tbl_text = texto_de(child)
            if 'competencias' in tbl_text.lower():
                for row in child.findall('.//' + W + 'tr')[1:]:
                    cells = row.findall(W + 'tc')
                    if len(cells) >= 2:
                        nombre_c = texto_de(cells[0]).rstrip('\xa0\u202f').strip()
                        nota_c   = texto_de(cells[1]).strip()
                        if nombre_c and nota_c:
                            competencias_lineas.append(f"{nombre_c}: {nota_c}")
                break
    competencias_txt = "\n".join(competencias_lineas)

    # Comentarios
    idx_consideraciones = next((i for i, t in enumerate(textos)
                                if t.strip().rstrip("\xa0 :") == "Consideraciones"), None)
    comentarios_texto = []
    capturando = False
    for t in textos:
        if t.startswith("Menciona un aspecto") or t.startswith("¿Tienes algún comentario"):
            capturando = True
            comentarios_texto.append(f"\n— {t} —")
            continue
        if t.strip().rstrip("\xa0 :") == "Consideraciones":
            break
        if capturando and t:
            comentarios_texto.append(t)

    return {
        "portada":    portada,
        "competencias": competencias_txt,
        "comentarios": "\n".join(comentarios_texto),
        "tiene_consideraciones_idx": idx_consideraciones is not None,
    }




def _primer_nombre(nombre_completo: str) -> str:
    """
    Extrae el primer nombre del docente.
    En el informe Word el nombre viene como: NOMBRE1 NOMBRE2 APELLIDO1 APELLIDO2
    (ej: "Juan Carlos Muñoz Mora") → el primer nombre es siempre partes[0].
    """
    partes = nombre_completo.strip().split()
    return partes[0].capitalize() if partes else "Docente"


def _detectar_genero(nombre_completo: str) -> str:
    """
    Devuelve 'M' o 'F' según terminación del primer nombre.
    Heurística simple: nombres terminados en 'a' (excepto excepciones) → femenino.
    """
    primer = _primer_nombre(nombre_completo).lower().rstrip(".")
    masculinos = {"luca", "nicola", "andrea", "bautista", "elias", "tobias",
                  "matias", "jeremias", "ezequias", "isaias", "josua", "josias",
                  "elia", "garcia"}
    femeninos_exc = set()
    if primer in masculinos:
        return "M"
    if primer.endswith("a") or primer.endswith("e"):
        return "F"
    return "M"


def _tratamiento(nombre_completo: str) -> tuple[str, str, str]:
    """
    Retorna (tratamiento, primer_nombre, genero)
    Ej: ("Profesora", "Laura", "F")
    """
    genero = _detectar_genero(nombre_completo)
    trat   = "Profesora" if genero == "F" else "Profesor"
    primer = _primer_nombre(nombre_completo)
    return trat, primer, genero

def llamar_github_models(token: str, prompt_sistema: str, prompt_usuario: str,
                          max_tokens: int = 1500, temperature: float = 0.4) -> str:
    """
    Llama a GitHub Models (chat completions) con gpt-4o y devuelve el texto de la respuesta.
    Lanza una excepción con mensaje claro si falla (token inválido, rate limit, etc.)
    """
    payload = {
        "model": GITHUB_MODELS_MODEL,
        "messages": [
            {"role": "system", "content": prompt_sistema},
            {"role": "user", "content": prompt_usuario},
        ],
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    data = _json.dumps(payload).encode("utf-8")
    req = _urlreq.Request(GITHUB_MODELS_URL, data=data, method="POST", headers={
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "Accept": "application/vnd.github+json",
    })
    try:
        with _urlreq.urlopen(req, timeout=90) as r:
            resultado = _json.loads(r.read())
        return resultado["choices"][0]["message"]["content"].strip()
    except _urlerr.HTTPError as e:
        cuerpo = e.read().decode(errors="replace")
        if e.code == 401:
            raise RuntimeError("Token de GitHub Models inválido o sin el permiso 'models: read'.") from e
        elif e.code == 429:
            raise RuntimeError("Se alcanzó el límite de solicitudes del tier gratuito (15/min o 150/día). "
                               "Espera un momento e inténtalo de nuevo.") from e
        else:
            raise RuntimeError(f"Error de GitHub Models ({e.code}): {cuerpo[:300]}") from e


def generar_consideraciones_ia(token: str, info_informe: dict, contexto_docs: str,
                                instruccion_usuaria: str = "",
                                nombre_docente: str = "") -> str:
    """Genera la retroalimentación formativa usando el prompt institucional EXA."""

    # Detectar primer nombre y tratamiento para personalizar el texto generado
    if nombre_docente:
        _trat, _pnombre, _genero = _tratamiento(nombre_docente)
        _el_la = "la" if _genero == "F" else "el"
    else:
        _trat, _pnombre, _el_la = "Profesor/a", "docente", "el/la"

    prompt_sistema = (
        "Actúa como un Diseñador Instruccional experto del Centro para la Excelencia en el "
        "Aprendizaje (EXA) de la Universidad EAFIT. Tu objetivo es analizar los resultados de "
        "la evaluación docente de un curso virtual o híbrido y generar una retroalimentación "
        "formativa, estratégica y empática, basada en el protocolo institucional, dirigiéndote "
        "en todo momento al docente de manera directa (usando el \"tú\" de forma cercana pero "
        "profesional). Al final, construirás una ruta de formación personalizada usando "
        "exclusivamente el catálogo de Aprende+ que se detalla más adelante.\n\n"

        "INSTRUCCIONES DE TAREA:\n"
        "Analiza los datos bajo los principios de feedforward y evaluación integral. "
        "Genera un informe formativo con las siguientes secciones:\n\n"

        "1. Panorama general y fortalezas (máx. 60 palabras — un solo párrafo)\n"
        "Redacta un único párrafo narrativo que abra con un reconocimiento cordial y concreto "
        "de las principales fortalezas, basadas exclusivamente en los comentarios cualitativos "
        "de los estudiantes — no menciones la tasa de respuesta ni datos cuantitativos. Luego, "
        "en continuidad fluida, ofrece una síntesis interpretativa del panorama general del "
        "desempeño. Usa un tono profesional, empático y de acompañamiento — nunca de juicio. "
        "Evita viñetas; todo debe fluir como prosa.\n\n"

        "2. Consideraciones y acciones de mejora\n"
        "Presenta entre 2 y 3 áreas de crecimiento. Para cada una, integra en un mismo bloque: "
        "la consideración (redactada con verbos como Revisar, Ajustar, Fomentar, Fortalecer, "
        "Evaluar, Promover, Regular, Realizar), la evidencia que la sustenta (comentario "
        "específico de los estudiantes) y la acción SMART sugerida (específica, medible, "
        "alcanzable, relevante y temporal para la siguiente cohorte). Agrupa ideas similares, "
        "evita repeticiones y mantén el lenguaje en clave de \"áreas de crecimiento\", no de "
        "\"deficiencias\".\n\n"

        "3. Ruta de formación personalizada\n"
        "Con base en las áreas de crecimiento identificadas, diseña una ruta de formación de "
        "1 a 2 pasos, ordenados de mayor a menor prioridad. Para cada paso indica: el recurso "
        "de Aprende+ recomendado (nombre exacto, enlace y una frase que explique por qué es "
        "relevante para este docente en particular).\n\n"

        "Usa ÚNICAMENTE los recursos del siguiente catálogo institucional:\n\n"
        "TRAYECTORIAS:\n"
        "① Diseño de Experiencias de Aprendizaje\n"
        "   → Competencia: diseñar y liderar secuencias didácticas con el Ciclo de Kolb.\n"
        "   → Úsala cuando: el docente necesita estructurar mejor sus actividades, diversificar "
        "metodologías o promover la participación y autonomía del estudiante.\n\n"
        "② Trayectoria Innovación\n"
        "   → Competencia: aplicar herramientas de innovación, creatividad y metodologías ágiles.\n"
        "   → Úsala cuando: el docente busca renovar su práctica con enfoques más creativos o ágiles.\n\n"
        "③ Trayectoria Inteligencia Artificial\n"
        "   → Competencia: integrar IA para optimizar diseño de recursos, automatizar procesos y "
        "mejorar eficiencia.\n"
        "   → Úsala cuando: el docente necesita mejorar su gestión de recursos digitales o quiere "
        "innovar con IA en el aula.\n\n"
        "CURSOS INDIVIDUALES:\n"
        "④ Diseño de Syllabus y Rúbricas para la Evaluación del Aprendizaje\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/143311\n"
        "   → Úsalo cuando: hay debilidad en claridad de criterios de evaluación o estructura del curso.\n\n"
        "⑤ Metodología del Ciclo de Kolb\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/130033\n"
        "   → Úsalo cuando: el docente requiere una introducción al aprendizaje experiencial antes "
        "de abordar la trayectoria completa.\n\n"
        "⑥ Evaluación del Aprendizaje\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/184752\n"
        "   → Úsalo cuando: se detecta debilidad en estrategias de evaluación, retroalimentación "
        "o seguimiento del aprendizaje.\n\n"
        "⑦ Aprendizaje Basado en Retos (ABR)\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/134289\n"
        "   → Úsalo cuando: el docente quiere incorporar metodologías activas que promuevan "
        "pensamiento crítico y colaboración.\n\n"
        "⑧ Aprendizaje Basado en Proyectos (ABP)\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/127810\n"
        "   → Úsalo cuando: el docente quiere que los estudiantes aprendan resolviendo situaciones "
        "reales mediante proyectos colaborativos.\n\n"
        "⑨ El Pacto Pedagógico\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/202260\n"
        "   → Úsalo cuando: se detectan problemas de comunicación, acuerdos de convivencia o "
        "compromiso entre docente y estudiantes.\n\n"
        "No inventes recursos ni enlaces. Si ningún recurso del catálogo se ajusta a una necesidad "
        "detectada, indícalo explícitamente.\n\n"

        "PRINCIPIOS DE REDACCIÓN:\n"
        "- Énfasis en \"áreas de crecimiento\", no en \"deficiencias\".\n"
        "- Feedforward: orienta hacia acciones futuras, no hacia errores del pasado.\n"
        "- Lo positivo siempre antes que las oportunidades de mejora.\n"
        "- Tono profesional, respetuoso y de acompañamiento (nunca de juicio).\n"
        "- Coherencia, buena gramática y puntuación en todo el texto.\n"
        "- Sin redundancias ni extensiones innecesarias.\n"
        "- Usa SOLO los recursos del catálogo proporcionado; no inventes ni agregues otros.\n"
        f"- Dirígete al docente usando su nombre en todas las menciones, con el formato "
        f"\"profesor/a + nombre propio\" (ejemplo: profesor {_pnombre} o profesora {_pnombre} "
        f"según corresponda). Su nombre es {_pnombre} y el tratamiento es {_trat}.\n"
        "- Los títulos de cada sección van en minúscula (salvo la primera letra) y en negrita.\n"
        "- No menciones la tasa de respuesta ni datos cuantitativos; basa el análisis "
        "exclusivamente en los comentarios cualitativos de los estudiantes.\n\n"

        "=== EJEMPLO DE SALIDA ESPERADA ===\n"
        "A continuación un ejemplo real de cómo debe verse el resultado. Imita este tono, "
        "estructura, nivel de detalle y formato exactamente:\n\n"

        "--- INICIO DEL EJEMPLO ---\n"
        "**Panorama general y fortalezas**\n"
        "Profesora María Adelaida, algunos estudiantes reconocen que el curso está bien "
        "organizado, que las instrucciones de las actividades suelen ser claras, que cumples "
        "con los pactos pedagógicos y que eres responsable en la entrega de calificaciones. "
        "Esos elementos de orden y estructura son una base sobre la cual trabajar, y las "
        "oportunidades de crecimiento que se identifican a continuación —que en este grupo se "
        "expresan con mayor intensidad que en otros— apuntan hacia la comunicación, la equidad "
        "evaluativa y el acompañamiento al estudiante en el entorno virtual.\n\n"

        "**Consideraciones y acciones de mejora**\n"
        "Área 1: Fortalecer la comunicación, la escucha activa y la disponibilidad ante "
        "inquietudes evaluativas\n"
        "Revisar la forma en que se gestionan las solicitudes de retroalimentación y revisión "
        "de notas, asegurando que todos los estudiantes reciban respuesta oportuna y respetuosa, "
        "especialmente en un formato virtual donde ese canal es el único punto de contacto. "
        "Múltiples estudiantes describieron intentos reiterados de comunicación por correo y "
        "Teams sin obtener respuesta, y en algunos casos recibieron respuestas que percibieron "
        "como poco empáticas. Como acción SMART: establece desde la primera semana un protocolo "
        "de comunicación con canales definidos, tiempos máximos de respuesta de 48 horas "
        "hábiles, y un procedimiento claro para la revisión de calificaciones, publicado en la "
        "plataforma y parte del pacto pedagógico inicial.\n\n"
        "Área 2: Revisar la coherencia y transparencia de los criterios de evaluación, "
        "especialmente en el proyecto final\n"
        "Ajustar las rúbricas y las instrucciones de las actividades evaluativas para que los "
        "criterios sean claros, aplicados de manera consistente y comunicados con el mismo nivel "
        "de detalle a todos los grupos. Varios estudiantes señalaron que la retroalimentación "
        "del proyecto final fue desigual entre grupos, que algunas instrucciones resultaron "
        "confusas —particularmente en relación con el uso de IA y Turnitin— y que las "
        "calificaciones no siempre correspondían con los criterios establecidos. Como acción "
        "SMART: antes de la próxima cohorte, revisa la rúbrica del proyecto final para que cada "
        "criterio tenga indicadores observables y verificables, define por escrito las condiciones "
        "de uso de herramientas de IA, y establece un protocolo uniforme de retroalimentación "
        "para todos los grupos.\n\n"

        "**Ruta de formación personalizada**\n"
        "El Pacto Pedagógico\n"
        "https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/202260\n"
        "Para ti, profesora María Adelaida, este recurso es especialmente urgente en este grupo, "
        "donde la desconexión comunicativa y la falta de acuerdos explícitos generaron "
        "situaciones de alta tensión entre docente y estudiantes. Construir ese pacto desde el "
        "primer día —con canales, tiempos, procedimientos y compromisos mutuos— puede prevenir "
        "gran parte de las fricciones identificadas y restablecer la confianza en el proceso "
        "formativo.\n\n"
        "Diseño de Syllabus y Rúbricas para la Evaluación del Aprendizaje\n"
        "https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/143311\n"
        "Este recurso te dará herramientas concretas para diseñar instrumentos evaluativos más "
        "transparentes, coherentes y aplicados de forma uniforme. Dado que en este grupo las "
        "inconformidades con la evaluación del proyecto final fueron el detonante principal de "
        "conflictos, contar con rúbricas bien construidas y criterios comunicados desde el inicio "
        "será clave para fortalecer la percepción de justicia y claridad en el curso.\n"
        "--- FIN DEL EJEMPLO ---"
    )


    partes_usuario = [
        f"INFORMACIÓN DEL CURSO Y PROFESOR:\n{info_informe['portada']}\n",
        f"COMENTARIOS CUALITATIVOS DE ESTUDIANTES:\n{info_informe['comentarios'] or '(sin comentarios registrados)'}\n",
    ]
    if contexto_docs:
        partes_usuario.append(f"DOCUMENTOS DE REFERENCIA INSTITUCIONALES:\n{contexto_docs}\n")

    partes_usuario.append(
        f"Con base en los comentarios cualitativos anteriores, genera el informe formativo. "
        f"Dirígete al docente como \"{_trat} {_pnombre}\". No menciones datos cuantitativos."
    )

    # El prompt editable del usuario ES el system prompt — manda sobre todo lo demás
    system = instruccion_usuaria.strip() if instruccion_usuaria.strip() else prompt_sistema

    prompt_usuario = "\n".join(partes_usuario)
    if len(prompt_usuario) > MAX_CHARS_CONTEXTO:
        prompt_usuario = prompt_usuario[:MAX_CHARS_CONTEXTO] + "\n[...contexto recortado por límite de tokens...]"

    return llamar_github_models(token, system, prompt_usuario, max_tokens=1500)


def _crear_parrafo_consideracion(texto: str, fuente: str = "Calibri",
                                  tam_pt: int = 11) -> etree._Element:
    """
    Crea un párrafo Word nuevo con:
    - Texto justificado (jc = both)
    - Espaciado después de párrafo: 160 twips (~8 pt) para separar bloques
    - Sangría de primera línea: 720 twips (1.27 cm)
    - Fuente y tamaño heredables de la plantilla (Calibri 11 por defecto)
    - Sin viñetas ni listas
    """
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    XML_SP = "http://www.w3.org/XML/1998/namespace"

    p = etree.Element(f"{{{W_NS}}}p")

    # ── Propiedades de párrafo ──
    pPr = etree.SubElement(p, f"{{{W_NS}}}pPr")

    # Justificado
    jc = etree.SubElement(pPr, f"{{{W_NS}}}jc")
    jc.set(f"{{{W_NS}}}val", "both")

    # Espaciado: 0 antes, 160 después (≈8 pt de separación entre párrafos)
    spacing = etree.SubElement(pPr, f"{{{W_NS}}}spacing")
    spacing.set(f"{{{W_NS}}}before", "0")
    spacing.set(f"{{{W_NS}}}after", "160")
    spacing.set(f"{{{W_NS}}}line", "276")      # interlineado 1.15
    spacing.set(f"{{{W_NS}}}lineRule", "auto")

    # Sangría primera línea (720 twips = 1.27 cm)
    ind = etree.SubElement(pPr, f"{{{W_NS}}}ind")
    ind.set(f"{{{W_NS}}}firstLine", "720")

    # ── Run con el texto ──
    r = etree.SubElement(p, f"{{{W_NS}}}r")

    # Propiedades del run: fuente y tamaño
    rPr = etree.SubElement(r, f"{{{W_NS}}}rPr")
    rFonts = etree.SubElement(rPr, f"{{{W_NS}}}rFonts")
    rFonts.set(f"{{{W_NS}}}ascii", fuente)
    rFonts.set(f"{{{W_NS}}}hAnsi", fuente)
    sz = etree.SubElement(rPr, f"{{{W_NS}}}sz")
    sz.set(f"{{{W_NS}}}val", str(tam_pt * 2))   # Word usa half-points
    szCs = etree.SubElement(rPr, f"{{{W_NS}}}szCs")
    szCs.set(f"{{{W_NS}}}val", str(tam_pt * 2))

    t_elem = etree.SubElement(r, f"{{{W_NS}}}t")
    t_elem.text = texto
    t_elem.set(f"{{{XML_SP}}}space", "preserve")

    return p


def insertar_consideraciones_en_docx(docx_bytes: bytes, texto_consideraciones: str) -> bytes:
    """
    Localiza el título 'Consideraciones' en el documento y reemplaza los párrafos
    vacíos que le siguen con párrafos justificados bien formateados.
    Cada bloque separado por línea en blanco se convierte en un párrafo independiente.
    """
    tree = etree.fromstring(
        zipfile.ZipFile(io.BytesIO(docx_bytes)).read('word/document.xml')
    )
    body = tree.find(W + 'body')
    children = list(body)

    def texto_de(elem):
        return "".join(t.text or "" for t in elem.iter(W + 't')).strip().rstrip("\xa0 ")

    idx_titulo = next((i for i, c in enumerate(children)
                       if texto_de(c).strip().rstrip(": ") == "Consideraciones"), None)
    if idx_titulo is None:
        raise RuntimeError("No se encontró la sección 'Consideraciones' en este informe. "
                           "Verifica que el archivo subido sea un informe generado por esta app.")

    # Buscar hasta dónde van los huecos (vacíos O placeholders) después del título
    idx_siguiente_con_texto = None
    for j in range(idx_titulo + 1, len(children)):
        t = texto_de(children[j])
        if t and '{ Comentario' not in t and '{Comentario' not in t:
            idx_siguiente_con_texto = j
            break
    if idx_siguiente_con_texto is None:
        idx_siguiente_con_texto = len(children)

    huecos = children[idx_titulo + 1: idx_siguiente_con_texto]
    insert_pos = (list(body).index(huecos[0]) if huecos
                  else list(body).index(children[idx_titulo]) + 1)

    # Eliminar huecos vacíos existentes
    for h in huecos:
        body.remove(h)

    # Partir el texto en párrafos (doble salto de línea = separador de bloque)
    parrafos_nuevos = [p.strip() for p in re.split(r"\n\s*\n", texto_consideraciones) if p.strip()]
    if not parrafos_nuevos:
        parrafos_nuevos = [texto_consideraciones.strip()]

    # Insertar cada párrafo correctamente formateado
    for k, parrafo in enumerate(parrafos_nuevos):
        nuevo_p = _crear_parrafo_consideracion(parrafo)
        body.insert(insert_pos + k, nuevo_p)

    new_xml = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
    out_buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zin:
        with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                zout.writestr(item, new_xml if item.filename == 'word/document.xml'
                              else zin.read(item.filename))
    return out_buf.getvalue()


# ─── INTERFAZ STREAMLIT ────────────────────────────────────────────────────

st.set_page_config(
    page_title="Generador de Informes — EAFIT",
    page_icon="📋",
    layout="centered",
)

# ── Colores oficiales EAFIT ──
# Azul:     #004B85   Amarillo: #FFB903
# Negro:    #000000   Superficie: #0C0C0E  Borde: #1C1C22

st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

  /* ── BASE ── */
  html, body, .stApp {
    background-color: #000000 !important;
    color: #E8EAF0 !important;
    font-family: 'Inter', sans-serif !important;
  }
  .block-container {
    padding-top: 0 !important;
    padding-bottom: 2rem !important;
    max-width: 720px !important;
  }

  /* ── HEADER ── */
  .exa-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 1.6rem 0 1.4rem 0;
    border-bottom: 1px solid #1C1C22;
    margin-bottom: 2rem;
  }
  .exa-header-left {
    display: flex;
    align-items: center;
    gap: 1rem;
  }
  .exa-logo img {
    height: 38px;
    width: auto;
    display: block;
  }
  .exa-divider-v {
    width: 1px;
    height: 36px;
    background: #1C1C22;
  }
  .exa-title-block {}
  .exa-title {
    font-size: 1.15rem;
    font-weight: 700;
    color: #FFFFFF;
    line-height: 1.25;
    letter-spacing: -0.01em;
    margin: 0;
  }
  .exa-subtitle {
    font-size: 0.76rem;
    color: #4A5068;
    margin: 3px 0 0 0;
    letter-spacing: 0.04em;
    text-transform: uppercase;
  }
  .exa-badge {
    font-size: 0.7rem;
    font-weight: 600;
    color: #000000;
    background: #FFB903;
    padding: 3px 10px;
    border-radius: 20px;
    letter-spacing: 0.05em;
    text-transform: uppercase;
  }

  /* ── ACCENT LINE ── */
  .exa-accent-line {
    height: 2px;
    background: linear-gradient(90deg, #004B85 0%, #FFB903 60%, transparent 100%);
    margin-bottom: 2rem;
    border-radius: 2px;
  }

  /* ── CARDS ── */
  .card {
    background: #0C0C0E;
    border: 1px solid #1C1C22;
    border-radius: 12px;
    padding: 1.4rem 1.6rem;
    margin-bottom: 1.2rem;
  }
  .card-label {
    font-size: 0.68rem;
    font-weight: 700;
    color: #FFB903;
    text-transform: uppercase;
    letter-spacing: 0.12em;
    margin-bottom: 1rem;
    display: flex;
    align-items: center;
    gap: 6px;
  }
  .card-label::after {
    content: '';
    flex: 1;
    height: 1px;
    background: #1C1C22;
  }

  /* ── INPUTS ── */
  .stTextInput > div > div > input {
    background: #070709 !important;
    color: #E8EAF0 !important;
    border: 1px solid #1C1C22 !important;
    border-radius: 8px !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.88rem !important;
  }
  .stTextInput > div > div > input:focus {
    border-color: #004B85 !important;
    box-shadow: 0 0 0 3px rgba(0,75,133,0.2) !important;
  }
  .stTextInput label {
    color: #9399A8 !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
  }

  /* ── FILE UPLOADER ── */
  .stFileUploader > div {
    background: #070709 !important;
    border: 1.5px dashed #1C1C22 !important;
    border-radius: 10px !important;
    transition: border-color 0.2s;
  }
  .stFileUploader > div:hover {
    border-color: #004B85 !important;
  }
  .stFileUploader label {
    color: #9399A8 !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
  }
  /* Upload icon/text color */
  .stFileUploader [data-testid="stFileUploaderDropzone"] p,
  .stFileUploader [data-testid="stFileUploaderDropzone"] span {
    color: #4A5068 !important;
  }

  /* ── BUTTON PRIMARY ── */
  .stButton > button {
    background: #004B85 !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    padding: 0.55rem 1.6rem !important;
    letter-spacing: 0.01em;
    transition: background 0.15s, box-shadow 0.15s !important;
  }
  .stButton > button:hover {
    background: #005FA8 !important;
    box-shadow: 0 0 0 3px rgba(0,75,133,0.25) !important;
  }
  .stButton > button:active { background: #003D6E !important; }

  /* ── BUTTON DOWNLOAD ── */
  .stDownloadButton > button {
    background: #FFB903 !important;
    color: #000000 !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    padding: 0.55rem 1.6rem !important;
    transition: background 0.15s, box-shadow 0.15s !important;
  }
  .stDownloadButton > button:hover {
    background: #FFC72C !important;
    box-shadow: 0 0 0 3px rgba(255,185,3,0.25) !important;
  }

  /* ── PROGRESS BAR ── */
  .stProgress > div > div {
    background: #1C1C22 !important;
    border-radius: 4px !important;
  }
  .stProgress > div > div > div {
    background: linear-gradient(90deg, #004B85, #FFB903) !important;
    border-radius: 4px !important;
  }

  /* ── ALERTS ── */
  .stAlert {
    background: #0C0C0E !important;
    border-radius: 8px !important;
    border-left-width: 3px !important;
  }
  [data-testid="stAlert"][kind="info"] {
    border-color: #004B85 !important;
  }
  [data-testid="stAlert"][kind="success"] {
    border-color: #16A34A !important;
  }
  [data-testid="stAlert"][kind="error"] {
    border-color: #DC2626 !important;
  }

  /* ── SPINNER ── */
  .stSpinner > div { border-top-color: #FFB903 !important; }

  /* ── PREVIEW TABLE ── */
  .preview-table {
    width: 100%;
    border-collapse: collapse;
    font-size: 0.82rem;
    margin-top: 0.4rem;
  }
  .preview-table thead tr {
    border-bottom: 1px solid #004B85;
  }
  .preview-table th {
    background: transparent;
    color: #4A5068;
    padding: 6px 10px;
    text-align: left;
    font-weight: 600;
    font-size: 0.68rem;
    text-transform: uppercase;
    letter-spacing: 0.1em;
  }
  .preview-table td {
    padding: 9px 10px;
    color: #C8CAD4;
    border-bottom: 1px solid #111116;
    vertical-align: middle;
  }
  .preview-table tr:last-child td { border-bottom: none; }
  .preview-table tr:hover td { background: #0F0F13; }
  .preview-table .name-cell { color: #FFFFFF; font-weight: 600; }
  .preview-table .file-cell { color: #4A5068; font-family: 'Courier New', monospace; font-size: 0.74rem; }
  .badge-ciclo {
    display: inline-block;
    background: #111116;
    color: #FFB903;
    border: 1px solid #1C1C22;
    border-radius: 4px;
    padding: 1px 7px;
    font-size: 0.72rem;
    font-weight: 600;
  }

  /* ── HIDE STREAMLIT CHROME ── */
  #MainMenu, footer, header { visibility: hidden; }
  [data-testid="stDecoration"] { display: none; }
</style>
""", unsafe_allow_html=True)

# ── Header ──
st.markdown("""
<div class="exa-header">
  <div class="exa-header-left">
    <div class="exa-logo">
      <img src="https://www.eafit.edu.co/sites/default/files/2024-07/logo_EAFIT_blanco.svg"
           alt="EAFIT" />
    </div>
    <div class="exa-divider-v"></div>
    <div class="exa-title-block">
      <div class="exa-title">Informes de evaluación docente</div>
      <div class="exa-subtitle">Centro para la Excelencia en el Aprendizaje · EXA</div>
    </div>
  </div>
  <div class="exa-badge">EXA</div>
</div>
<div class="exa-accent-line"></div>
""", unsafe_allow_html=True)

# ── Cargar plantilla desde el repositorio ──
_PLANTILLA_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Plantilla.docx")

@st.cache_data
def _cargar_plantilla():
    if os.path.isfile(_PLANTILLA_PATH):
        with open(_PLANTILLA_PATH, "rb") as f:
            return f.read()
    return None

plantilla_bytes = _cargar_plantilla()

if plantilla_bytes is None:
    st.error("⚠️ No se encontró **Plantilla.docx** en el repositorio. "
             "Asegúrate de subir ese archivo a GitHub junto con `app.py`.")
    st.stop()

# ── Menú de navegación (barra en el contenido principal, siempre visible) ──
if "pagina_actual" not in st.session_state:
    st.session_state.pagina_actual = "generar"

col_nav1, col_nav2 = st.columns(2)
with col_nav1:
    if st.button("📋  Generar informe", use_container_width=True,
                 type=("primary" if st.session_state.pagina_actual == "generar" else "secondary")):
        st.session_state.pagina_actual = "generar"
        st.rerun()
with col_nav2:
    if st.button("✨  Consideraciones", use_container_width=True,
                 type=("primary" if st.session_state.pagina_actual == "consideraciones" else "secondary")):
        st.session_state.pagina_actual = "consideraciones"
        st.rerun()

PAGINA = st.session_state.pagina_actual



# ── Base de datos de evaluaciones (en disco, no se carga en RAM) ──
_DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "evaluaciones.xlsx")

# ── Entrada: Catálogo y Nº Clase ──
if PAGINA == "generar":
    if not os.path.isfile(_DB_PATH):
        st.error("⚠️ No se encontró **evaluaciones.xlsx** en el repositorio. "
                 "Asegúrate de subirlo a GitHub junto con `app.py`.")
        st.stop()

    st.markdown('<div class="card"><div class="card-label">🔍 Buscar clase</div>', unsafe_allow_html=True)

    input_codigo = st.text_input(
        "Catálogo – Nº de clase",
        placeholder="Ej: OG2117-5890",
        help="Ingresa el código del catálogo seguido de un guion y el número de clase (ej: OG2117-5890)"
    )

    buscar = st.button("🔎 Buscar y previsualizar")
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Estado de sesión para profesores encontrados ──
    if "profesores" not in st.session_state:
        st.session_state.profesores = {}

    profesores = st.session_state.profesores

    if buscar:
        codigo = input_codigo.strip()
        if not codigo or "-" not in codigo:
            st.warning("Ingresa el código en el formato **CATÁLOGO-CLASE**, por ejemplo: `OG2117-5890`.")
        else:
            # Separar por el último guion para admitir catálogos con letras+números
            partes = codigo.rsplit("-", 1)
            if len(partes) != 2 or not partes[0].strip() or not partes[1].strip():
                st.warning("Formato inválido. Usa **CATÁLOGO-CLASE**, por ejemplo: `OG2117-5890`.")
            else:
                input_catalogo = partes[0].strip()
                input_clase    = partes[1].strip()
                try:
                    with st.spinner("Buscando en la base de datos…"):
                        resultado = leer_excel(
                            archivo_path=_DB_PATH,
                            filtro_catalogo=input_catalogo,
                            filtro_clase=input_clase
                        )
                    if not resultado:
                        st.error(f"No se encontraron registros para **{input_catalogo.upper()}-{input_clase}**. "
                                 "Verifica el catálogo y número de clase.")
                        st.session_state.profesores = {}
                    else:
                        st.session_state.profesores = resultado
                        profesores = resultado
                except Exception as e:
                    st.error(f"Error al leer la base de datos: {e}")

    # ── Preview de resultados ──
    if profesores:
        try:
            st.markdown('<div class="card"><div class="card-label">👥 Profesores encontrados</div>',
                        unsafe_allow_html=True)

            filas_html = ""
            for nombre, datos in profesores.items():
                info = datos["info"]
                nf   = nombre_archivo_defecto(datos, nombre)
                filas_html += f"""
                <tr>
                  <td class="name-cell">{nombre.title()}</td>
                  <td>{info.get('curso','—')}</td>
                  <td>{info.get('escuela','—')}</td>
                  <td><span class="badge-ciclo">{info.get('ciclo','—')}</span></td>
                  <td class="file-cell">{nf}.docx</td>
                </tr>"""

            st.markdown(f"""
            <table class="preview-table">
              <thead>
                <tr>
                  <th>Profesor</th><th>Curso</th><th>Escuela</th>
                  <th>Semestre</th><th>Nombre del archivo</th>
                </tr>
              </thead>
              <tbody>{filas_html}</tbody>
            </table>
            """, unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        except Exception as e:
            st.error(f"Error al mostrar resultados: {e}")

    # ── Sección: Generar informes ──
    if profesores:
        st.markdown('<div class="card"><div class="card-label">⚙️ Generar informes</div>',
                    unsafe_allow_html=True)

        total = len(profesores)
        es_uno = total == 1

        nombre_custom = None
        if es_uno:
            nombre_prof, datos_prof = next(iter(profesores.items()))
            defecto = nombre_archivo_defecto(datos_prof, nombre_prof)
            nombre_custom = st.text_input(
                "Nombre del archivo (editable)",
                value=defecto,
                help="Puedes cambiar el nombre antes de generar"
            )

        st.markdown('</div>', unsafe_allow_html=True)

        if st.button(f"Generar {'informe' if es_uno else str(total) + ' informes'}"):
            errores = []
            archivos_generados = {}

            barra = st.progress(0, text="Generando informes…")
            for i, (nombre, datos) in enumerate(profesores.items()):
                try:
                    nf = nombre_custom if (es_uno and nombre_custom) else None
                    docx_bytes, nombre_arch = generar_informe_bytes(
                        nombre, datos, plantilla_bytes, nombre_archivo=nf
                    )
                    archivos_generados[nombre_arch + ".docx"] = docx_bytes
                except Exception as e:
                    errores.append(f"{nombre}: {e}")
                barra.progress((i + 1) / total,
                               text=f"Procesando {i+1} de {total}…")

            barra.empty()

            if errores:
                for err in errores:
                    st.error(f"❌ {err}")

            if archivos_generados:
                if len(archivos_generados) == 1:
                    nombre_arch, docx_bytes = next(iter(archivos_generados.items()))
                    st.session_state.ultimo_informe_generado = {
                        "nombre": nombre_arch,
                        "bytes": docx_bytes,
                    }
                    # Limpiar cache de consideraciones para que se refresque con el nuevo informe
                    for _k in ["_nombre_docente_cache_key", "_nombre_docente_cache",
                               "_consideraciones_informe_key", "consideraciones_texto",
                               "consideraciones_docx_bytes", "consideraciones_nombre_archivo"]:
                        st.session_state.pop(_k, None)
                    st.success(f"✅ Informe generado: **{nombre_arch}**")
                    st.download_button(
                        label="⬇️  Descargar informe",
                        data=docx_bytes,
                        file_name=nombre_arch,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                else:
                    zip_buf = io.BytesIO()
                    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for nombre_arch, docx_bytes in archivos_generados.items():
                            zf.writestr(nombre_arch, docx_bytes)

                    st.success(f"✅ {len(archivos_generados)} informes generados correctamente.")
                    st.download_button(
                        label="⬇️  Descargar todos los informes (.zip)",
                        data=zip_buf.getvalue(),
                        file_name="Informes_EAFIT.zip",
                        mime="application/zip",
                    )

    elif buscar and not profesores:
        pass  # El error ya se mostró arriba
    else:
        st.info("🔍 Ingresa el código en formato **CATÁLOGO-CLASE** (ej: `OG2117-5890`) para comenzar.")

# ── Sección: Actualizar base de datos (al final) ──
import base64
import urllib.request
import urllib.error
import json as _json

_GH_REPO  = "Sof-Saos/informes-evaluaciondocente-"   # usuario/repo
_GH_FILE  = "evaluaciones.xlsx"                        # ruta dentro del repo
_GH_TOKEN = st.secrets.get("GITHUB_TOKEN", "")        # secreto en Streamlit Cloud
_GH_DOCS_FOLDER = "docs_guia"                          # carpeta de documentos guía persistentes

def _gh_get_sha(token: str) -> str | None:
    """Obtiene el SHA actual del archivo evaluaciones.xlsx en GitHub."""
    url = f"https://api.github.com/repos/{_GH_REPO}/contents/{_GH_FILE}"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    })
    try:
        with urllib.request.urlopen(req, timeout=10) as r:
            return _json.loads(r.read())["sha"]
    except urllib.error.HTTPError as e:
        if e.code == 404:
            return None
        raise

def _gh_get_file_sha(token: str, path: str) -> str | None:
    """Obtiene el SHA de cualquier archivo en el repo (necesario para actualizarlo)."""
    url = f"https://api.github.com/repos/{_GH_REPO}/contents/{path}"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    })
    try:
        with urllib.request.urlopen(req, timeout=10) as r:
            return _json.loads(r.read())["sha"]
    except urllib.error.HTTPError as e:
        if e.code == 404:
            return None
        raise

def _gh_upload_doc_guia(token: str, nombre: str, contenido: bytes) -> bool:
    """Sube un documento guía a docs_guia/ en el repo. Si ya existe, lo actualiza."""
    import base64
    # Sanitizar: reemplazar espacios y caracteres problemáticos en la URL
    nombre_safe = re.sub(r'[^A-Za-z0-9._\-]', '_', nombre)
    path = f"{_GH_DOCS_FOLDER}/{nombre_safe}"
    url  = f"https://api.github.com/repos/{_GH_REPO}/contents/{path}"
    sha  = _gh_get_file_sha(token, path)   # None si no existe aún
    payload = {
        "message": f"Agregar/actualizar documento guía: {nombre}",
        "content": base64.b64encode(contenido).decode(),
        "branch": "main",
    }
    if sha:
        payload["sha"] = sha
    data = _json.dumps(payload).encode()
    req = urllib.request.Request(url, data=data, method="PUT", headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "Content-Type": "application/json",
        "X-GitHub-Api-Version": "2022-11-28",
    })
    with urllib.request.urlopen(req, timeout=60) as r:
        return r.status in (200, 201)

def _gh_delete_doc_guia(token: str, nombre: str) -> bool:
    """Elimina un documento guía de docs_guia/ en el repo."""
    path = f"{_GH_DOCS_FOLDER}/{nombre}"
    sha  = _gh_get_file_sha(token, path)
    if not sha:
        return True   # ya no existe
    url  = f"https://api.github.com/repos/{_GH_REPO}/contents/{path}"
    payload = {
        "message": f"Eliminar documento guía: {nombre}",
        "sha": sha,
        "branch": "main",
    }
    data = _json.dumps(payload).encode()
    req = urllib.request.Request(url, data=data, method="DELETE", headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "Content-Type": "application/json",
        "X-GitHub-Api-Version": "2022-11-28",
    })
    with urllib.request.urlopen(req, timeout=30) as r:
        return r.status in (200, 201)

_GH_PROMPT_FILE = "prompt_consideraciones.txt"   # archivo donde se guarda el prompt

def _gh_cargar_prompt(token: str) -> str | None:
    """Carga el prompt guardado desde GitHub. Devuelve None si no existe aún."""
    import base64
    url = f"https://api.github.com/repos/{_GH_REPO}/contents/{_GH_PROMPT_FILE}"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    })
    try:
        with urllib.request.urlopen(req, timeout=10) as r:
            data = _json.loads(r.read())
            return base64.b64decode(data["content"]).decode("utf-8")
    except urllib.error.HTTPError as e:
        if e.code == 404:
            return None
        raise

def _gh_guardar_prompt(token: str, prompt_texto: str) -> bool:
    """Guarda el prompt en GitHub (crea o actualiza prompt_consideraciones.txt)."""
    import base64
    path = _GH_PROMPT_FILE
    sha  = _gh_get_file_sha(token, path)
    payload = {
        "message": "Actualizar prompt de consideraciones",
        "content": base64.b64encode(prompt_texto.encode("utf-8")).decode(),
        "branch":  "main",
    }
    if sha:
        payload["sha"] = sha
    url  = f"https://api.github.com/repos/{_GH_REPO}/contents/{path}"
    data = _json.dumps(payload).encode()
    req  = urllib.request.Request(url, data=data, method="PUT", headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "Content-Type": "application/json",
        "X-GitHub-Api-Version": "2022-11-28",
    })
    with urllib.request.urlopen(req, timeout=30) as r:
        return r.status in (200, 201)

@st.cache_data(ttl=300, show_spinner=False)
def _gh_cargar_prompt_cached(token: str) -> str | None:
    """Versión cacheada (5 min) de _gh_cargar_prompt para no hammear la API."""
    return _gh_cargar_prompt(token)

@st.cache_data(ttl=60, show_spinner=False)
def _gh_listar_docs_guia(token: str) -> list[dict]:
    """
    Lista los documentos guardados en docs_guia/ del repo.
    Devuelve lista de dicts: [{nombre, download_url, sha, size}].
    Cacheado 60s para no hammear la API en cada rerun de Streamlit.
    """
    url = f"https://api.github.com/repos/{_GH_REPO}/contents/{_GH_DOCS_FOLDER}"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    })
    try:
        with urllib.request.urlopen(req, timeout=15) as r:
            items = _json.loads(r.read())
            return [
                {"nombre": i["name"], "url": i["download_url"],
                 "sha": i["sha"], "size": i.get("size", 0)}
                for i in items if i["type"] == "file"
            ]
    except urllib.error.HTTPError as e:
        if e.code == 404:
            return []   # carpeta no existe aún
        raise


    """
    Lista los documentos guardados en docs_guia/ del repo.
    Devuelve lista de dicts: [{nombre, download_url, sha, size}].
    Cacheado 60s para no hammear la API en cada rerun de Streamlit.
    """
    url = f"https://api.github.com/repos/{_GH_REPO}/contents/{_GH_DOCS_FOLDER}"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
    })
    try:
        with urllib.request.urlopen(req, timeout=15) as r:
            items = _json.loads(r.read())
            return [
                {"nombre": i["name"], "url": i["download_url"],
                 "sha": i["sha"], "size": i.get("size", 0)}
                for i in items if i["type"] == "file"
            ]
    except urllib.error.HTTPError as e:
        if e.code == 404:
            return []   # carpeta no existe aún
        raise

def _gh_descargar_doc(url: str, token: str) -> bytes:
    """Descarga el contenido de un archivo desde su download_url de GitHub."""
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github.raw",
    })
    with urllib.request.urlopen(req, timeout=30) as r:
        return r.read()

def _comprimir_excel(archivo_bytes: bytes) -> bytes:
    """
    Comprime un Excel sin perder datos:
    - Recarga con openpyxl y guarda con ZIP_DEFLATED máximo.
    - Elimina estilos innecesarios y vistas en caché.
    Retorna los bytes comprimidos (puede ser menor o igual al original).
    """
    import zipfile, io
    # Estrategia 1: re-empaquetar el ZIP del xlsx con compresión máxima
    input_zip  = zipfile.ZipFile(io.BytesIO(archivo_bytes), "r")
    out_buf    = io.BytesIO()
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as out_zip:
        for item in input_zip.infolist():
            data = input_zip.read(item.filename)
            out_zip.writestr(item, data)
    input_zip.close()
    compressed = out_buf.getvalue()

    # Devolver el más pequeño (siempre debería ser el comprimido)
    return compressed if len(compressed) < len(archivo_bytes) else archivo_bytes


def _gh_push_file(token: str, content_bytes: bytes, sha: str | None, mensaje: str, intentos: int = 2) -> bool:
    """Sube (o reemplaza) el archivo en GitHub mediante un commit.
    Reintenta una vez con el SHA más reciente si GitHub responde 409 (conflicto)."""
    url = f"https://api.github.com/repos/{_GH_REPO}/contents/{_GH_FILE}"
    for intento in range(intentos):
        payload = {
            "message": mensaje,
            "content": base64.b64encode(content_bytes).decode(),
            "branch": "main",
        }
        if sha:
            payload["sha"] = sha
        data = _json.dumps(payload).encode()
        req = urllib.request.Request(url, data=data, method="PUT", headers={
            "Authorization": f"Bearer {token}",
            "Accept": "application/vnd.github+json",
            "Content-Type": "application/json",
            "X-GitHub-Api-Version": "2022-11-28",
        })
        try:
            with urllib.request.urlopen(req, timeout=60) as r:
                return r.status in (200, 201)
        except urllib.error.HTTPError as e:
            if e.code == 409 and intento < intentos - 1:
                # Conflicto: el SHA quedó desactualizado, volvemos a consultarlo y reintentamos
                sha = _gh_get_sha(token)
                continue
            raise
    return False

if PAGINA == "generar":
    st.markdown("---")
    with st.expander("🔄 Actualizar base de datos de evaluación docente"):
        if not _GH_TOKEN:
            st.warning(
                "⚠️ No se encontró el secreto **GITHUB_TOKEN** en Streamlit Cloud. "
                "Agrégalo en *Settings → Secrets* para habilitar la actualización permanente."
            )
        else:
            st.markdown(
                "<small style='color:#4A5068'>Sube el Excel del nuevo semestre. "
                "El archivo se guardará directamente en GitHub — el cambio es <b>permanente</b> "
                "y la app se actualizará automáticamente en unos segundos.</small>",
                unsafe_allow_html=True
            )
            nuevo_excel = st.file_uploader(
                "Nuevo archivo de evaluaciones (.xlsx)",
                type=["xlsx"],
                key="uploader_db"
            )
            if nuevo_excel:
                if st.button("✅ Confirmar actualización de base de datos"):
                    try:
                        nuevos_bytes = nuevo_excel.getvalue()

                        # 1. Validar estructura del Excel
                        wb_test = openpyxl.load_workbook(io.BytesIO(nuevos_bytes), read_only=True)
                        ws_test = wb_test.active
                        headers_test = {}
                        for row in ws_test.iter_rows(min_row=FILA_ENCABEZADO, max_row=FILA_ENCABEZADO, values_only=True):
                            for i, val in enumerate(row):
                                if val: headers_test[str(val).strip()] = i
                            break
                        cols_requeridas = [COL_NOMBRE, COL_CATALOGO, COL_NCLASE, COL_COMPETENCIA, COL_NOTA_FINAL]
                        faltantes = [c for c in cols_requeridas if c not in headers_test]
                        tamano_mb = len(nuevos_bytes) / (1024 * 1024)
                        if faltantes:
                            st.error(f"El archivo no tiene las columnas requeridas: {', '.join(faltantes)}")
                        else:
                            # ── Compresión automática si el archivo es pesado ──
                            if tamano_mb > 20:
                                with st.spinner(f"El archivo pesa {tamano_mb:.1f} MB — comprimiendo automáticamente…"):
                                    nuevos_bytes_comp = _comprimir_excel(nuevos_bytes)
                                    tamano_comp_mb    = len(nuevos_bytes_comp) / (1024 * 1024)
                                ahorro = tamano_mb - tamano_comp_mb
                                if ahorro > 0.05:
                                    st.info(
                                        f"📦 Archivo comprimido: {tamano_mb:.1f} MB → **{tamano_comp_mb:.1f} MB** "
                                        f"(ahorro de {ahorro:.1f} MB, sin pérdida de datos)."
                                    )
                                    nuevos_bytes = nuevos_bytes_comp
                                    tamano_mb    = tamano_comp_mb
                                else:
                                    st.info(f"ℹ️ El archivo ya estaba bien comprimido ({tamano_mb:.1f} MB).")

                            if tamano_mb > 99:
                                st.error(
                                    f"❌ El archivo pesa **{tamano_mb:.1f} MB** incluso tras comprimirlo. "
                                    "La API de GitHub tiene un límite de ~100 MB. "
                                    "Elimina hojas o columnas que no se usen y vuelve a intentarlo."
                                )
                            else:
                                # 2. Subir a GitHub
                                with st.spinner("Subiendo a GitHub…"):
                                    sha_actual = _gh_get_sha(_GH_TOKEN)
                                    ciclo_val  = ws_test.cell(row=1, column=2).value or ""
                                    commit_msg = f"Actualizar evaluaciones.xlsx — ciclo {ciclo_val} ({len(nuevos_bytes)//1024} KB)"
                                    ok = _gh_push_file(_GH_TOKEN, nuevos_bytes, sha_actual, commit_msg)
                                if ok:
                                    st.cache_data.clear()
                                    st.success(
                                        f"✅ Base de datos actualizada en GitHub ({len(nuevos_bytes)//1024} KB). "
                                        "Streamlit Cloud redeployará la app en unos segundos con los datos nuevos."
                                    )
                                else:
                                    st.error("No se pudo subir el archivo a GitHub. Verifica el token y los permisos.")
                    except urllib.error.HTTPError as e:
                        body = e.read().decode(errors="replace")
                        st.error(f"Error de GitHub ({e.code}): {body}")
                    except Exception as e:
                        st.error(f"Error al actualizar: {e}")


# ═══════════════════════════════════════════════════════════════════════════
# PÁGINA: ALISTAMIENTO DE CONSIDERACIONES
# ═══════════════════════════════════════════════════════════════════════════

if PAGINA == "consideraciones":
    _GH_MODELS_TOKEN = st.secrets.get("GITHUB_MODELS_TOKEN", "")

    st.markdown(
        '<div class="card"><div class="card-label">✨ Alistamiento de Consideraciones</div>'
        '<p style="color:#A8ACBE;font-size:0.88rem;margin:0.3rem 0 0 0">'
        'Sube un informe ya generado por esta app y, opcionalmente, documentos guía '
        '(formaciones EXA, protocolos, lineamientos, trabajo de grado) para que la IA '
        'reescriba únicamente la sección <b>Consideraciones</b>. El resto del informe '
        'no se modifica.</p></div>',
        unsafe_allow_html=True
    )

    if not _GH_MODELS_TOKEN:
        st.warning(
            "⚠️ No se encontró el secreto **GITHUB_MODELS_TOKEN** en Streamlit Cloud. "
            "Este módulo necesita un token de GitHub con permiso **models: read** para "
            "funcionar. Ve a *Settings → Secrets* en Streamlit Cloud y agrégalo."
        )
        st.stop()

    # ── Paso 1: Informe base (detectado automáticamente o subido) ──
    st.markdown('<div class="card"><div class="card-label">1️⃣ Informe base (.docx)</div>',
                unsafe_allow_html=True)

    ultimo_informe = st.session_state.get("ultimo_informe_generado")

    usar_otro = st.checkbox(
        "Añadir consideraciones para otro informe (subir un .docx distinto)",
        value=(ultimo_informe is None),   # si no hay informe reciente, se activa solo
        key="usar_otro_informe"
    )

    informe_bytes_base  = None
    informe_nombre_base = None

    if not usar_otro and ultimo_informe is not None:
        # Detectar si el informe cambió respecto al que teníamos cacheado
        _informe_key = ultimo_informe.get("nombre", "")
        if st.session_state.get("_consideraciones_informe_key") != _informe_key:
            # Informe nuevo — limpiar consideraciones anteriores y forzar refresco
            st.session_state["_consideraciones_informe_key"] = _informe_key
            for k in ["consideraciones_texto", "consideraciones_docx_bytes",
                      "consideraciones_nombre_archivo", "_nombre_docente_cache"]:
                st.session_state.pop(k, None)
            st.rerun()

        st.success(f"✅ Documento generado encontrado: **{ultimo_informe['nombre']}**")
        informe_bytes_base  = ultimo_informe["bytes"]
        informe_nombre_base = ultimo_informe["nombre"]
    elif not usar_otro and ultimo_informe is None:
        st.info("Todavía no se ha generado ningún informe en esta sesión. Marca la casilla para subir uno manualmente.")
    else:
        informe_subido = st.file_uploader(
            "Sube el informe (.docx) generado por esta app",
            type=["docx"],
            key="uploader_informe_base"
        )
        if informe_subido is not None:
            informe_bytes_base  = informe_subido.getvalue()
            informe_nombre_base = informe_subido.name

    st.markdown('</div>', unsafe_allow_html=True)

    # ── Paso 2: Documentos guía (persistentes en GitHub + subida de nuevos) ──
    st.markdown('<div class="card"><div class="card-label">2️⃣ Documentos guía</div>',
                unsafe_allow_html=True)

    docs_seleccionados_nombres = []   # nombres de los docs persistentes a usar
    docs_nuevos_bytes = []            # [(nombre, bytes)] de archivos recién subidos

    if not _GH_TOKEN:
        st.warning("⚠️ Sin **GITHUB_TOKEN** configurado, los documentos no se pueden guardar permanentemente. "
                   "Puedes usarlos solo para esta sesión.")
        docs_guia_sesion = st.file_uploader(
            "Sube documentos guía (PDF, DOCX, TXT)",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            key="uploader_docs_guia_sesion"
        )
        docs_nuevos_bytes = [(d.name, d.getvalue()) for d in (docs_guia_sesion or [])]
    else:
        # ── Documentos ya guardados en GitHub ──
        try:
            docs_persistentes = _gh_listar_docs_guia(_GH_TOKEN)
        except Exception as e:
            docs_persistentes = []
            st.warning(f"No se pudo listar los documentos guardados: {e}")

        if docs_persistentes:
            st.markdown(
                "<small style='color:#A8ACBE'>Documentos guardados en el repositorio — "
                "marca los que quieres usar como referencia para la IA:</small>",
                unsafe_allow_html=True
            )
            col_check, col_del = st.columns([5, 1])
            for doc in docs_persistentes:
                size_kb = doc["size"] // 1024
                with col_check:
                    seleccionado = st.checkbox(
                        f"📄 {doc['nombre']} ({size_kb} KB)",
                        value=True,
                        key=f"doc_persistente_{doc['nombre']}"
                    )
                    if seleccionado:
                        docs_seleccionados_nombres.append(doc)
                with col_del:
                    st.markdown("<div style='margin-top:0.4rem'>", unsafe_allow_html=True)
                    if st.button("🗑️", key=f"del_{doc['nombre']}",
                                 help=f"Eliminar {doc['nombre']} permanentemente del repositorio"):
                        try:
                            with st.spinner(f"Eliminando {doc['nombre']}…"):
                                _gh_delete_doc_guia(_GH_TOKEN, doc["nombre"])
                                st.cache_data.clear()
                            st.success(f"✅ {doc['nombre']} eliminado.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error al eliminar: {e}")
                    st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.info("Todavía no hay documentos guía guardados. Sube los primeros abajo.")

        # ── Subir nuevos documentos (se guardan en GitHub automáticamente) ──
        st.markdown(
            "<small style='color:#A8ACBE;margin-top:0.8rem;display:block'>"
            "Sube nuevos documentos — quedarán guardados permanentemente en el repositorio "
            "y disponibles en todas las sesiones futuras:</small>",
            unsafe_allow_html=True
        )
        nuevos_subidos = st.file_uploader(
            "Formaciones EXA, protocolos, lineamientos, trabajo de grado…",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            key="uploader_docs_guia_nuevos"
        )

        # Guardar en session_state en cuanto se suben (antes de cualquier botón/rerun)
        if nuevos_subidos:
            st.session_state["_docs_pendientes"] = [
                (d.name, d.getvalue()) for d in nuevos_subidos
            ]

        docs_pendientes = st.session_state.get("_docs_pendientes", [])

        if docs_pendientes:
            # Los docs pendientes SIEMPRE están disponibles para la IA en esta sesión,
            # independientemente de si se guardan o no en GitHub.
            docs_nuevos_bytes = [(n, b) for n, b in docs_pendientes]

            col_save, col_info = st.columns([2, 3])
            with col_save:
                if st.button("💾 Guardar en el repositorio", key="btn_guardar_docs",
                             help="Los guarda permanentemente para todas las sesiones futuras"):
                    errores_subida = []
                    with st.spinner(f"Guardando {len(docs_pendientes)} documento(s)…"):
                        for nombre_doc, bytes_doc in docs_pendientes:
                            try:
                                _gh_upload_doc_guia(_GH_TOKEN, nombre_doc, bytes_doc)
                            except Exception as e:
                                errores_subida.append(f"{nombre_doc}: {e}")
                    if errores_subida:
                        for err in errores_subida:
                            st.error(f"❌ {err}")
                    else:
                        st.session_state.pop("_docs_pendientes", None)
                        st.cache_data.clear()
                        st.success(f"✅ {len(docs_pendientes)} documento(s) guardado(s) en el repositorio.")
                        st.rerun()
            with col_info:
                st.markdown(
                    f"<small style='color:#A8ACBE;line-height:2.2'>"
                    f"📂 {len(docs_pendientes)} doc(s) listos para usar en esta sesión. "
                    f"Guárdalos para que estén disponibles siempre.</small>",
                    unsafe_allow_html=True
                )

    st.markdown('</div>', unsafe_allow_html=True)

    # ── Paso 3: Prompt editable ──
    st.markdown('<div class="card"><div class="card-label">3️⃣ Prompt para la IA (editable)</div>',
                unsafe_allow_html=True)
    st.markdown(
        '<small style="color:#A8ACBE">Este es el prompt base que se le envía a la IA. '        'Puedes modificarlo, añadir instrucciones específicas o cambiar el enfoque antes de generar.</small>',
        unsafe_allow_html=True
    )

    # ── Detectar nombre y género del docente desde el informe subido ──
    # Extraer nombre del docente — solo si el informe cambió (se cachea en session_state)
    _cache_key = informe_nombre_base or ""
    if st.session_state.get("_nombre_docente_cache_key") != _cache_key:
        _nombre_docente_raw = ""
        if informe_bytes_base is not None:
            try:
                _info_tmp = extraer_texto_informe_actual(informe_bytes_base)
                for linea in _info_tmp.get("portada", "").splitlines():
                    linea_strip = linea.strip()
                    if re.search(r"[Nn]ombre\s+profe", linea_strip):
                        m = re.search(r"[Nn]ombre\s+profe[^:]*:\s*(.+)", linea_strip)
                        if m:
                            _nombre_docente_raw = m.group(1).strip()
                            break
                        m2 = re.search(r"[Nn]ombre\s+[Pp]rofesor[a]?\s+(.+)", linea_strip)
                        if m2:
                            _nombre_docente_raw = m2.group(1).strip()
                            break
                if not _nombre_docente_raw and ultimo_informe:
                    nombre_arch = ultimo_informe.get("nombre", "")
                    partes_arch = nombre_arch.replace(".docx", "").split("_")
                    if len(partes_arch) >= 4:
                        _nombre_docente_raw = partes_arch[-1]
            except Exception:
                pass
        st.session_state["_nombre_docente_cache_key"] = _cache_key
        st.session_state["_nombre_docente_cache"] = _nombre_docente_raw
    else:
        _nombre_docente_raw = st.session_state.get("_nombre_docente_cache", "")

    if _nombre_docente_raw:
        _trat, _pnombre, _genero = _tratamiento(_nombre_docente_raw)
        st.info(f"👤 Docente detectado/a: **{_trat} {_pnombre}** — la IA se dirigirá a {('ella' if _genero == 'F' else 'él')} por este nombre. *(nombre completo leído: {_nombre_docente_raw})*")

    # Prompt editable: solo las instrucciones institucionales (el saludo lo inserta la IA automáticamente)
    PROMPT_BASE_USUARIO = (
        "Actúa como un Diseñador Instruccional experto del Centro para la Excelencia en el "
        "Aprendizaje (EXA) de la Universidad EAFIT. Tu objetivo es analizar los resultados de "
        "la evaluación docente de un curso virtual o híbrido y generar una retroalimentación "
        "formativa, estratégica y empática, basada en el protocolo institucional, dirigiéndote "
        "en todo momento al docente de manera directa (usando el \"tú\" de forma cercana pero "
        "profesional). Al final, construirás una ruta de formación personalizada usando "
        "exclusivamente el catálogo de Aprende+ que se detalla más adelante.\n\n"
        "INSTRUCCIONES DE TAREA:\n"
        "Analiza los datos anteriores bajo los principios de feedforward y evaluación integral. "
        "Genera un informe formativo con las siguientes secciones:\n\n"
        "1. PANORAMA GENERAL Y FORTALEZAS (máx. 300 palabras — un solo párrafo)\n"
        "Redacta un único párrafo narrativo que cumpla las dos funciones a la vez: abre con un "
        "reconocimiento cordial y concreto de sus principales fortalezas, basadas en los "
        "comentarios de los estudiantes y los datos cuantitativos; luego, en continuidad fluida, "
        "ofrece una síntesis interpretativa del panorama general del desempeño. Destaca 2 de sus "
        "fortalezas y anuncia con lenguaje constructivo las 2-3 áreas que se abordarán a "
        "continuación. Usa un tono profesional, empático y de acompañamiento — nunca de juicio. "
        "Evita viñetas; todo debe fluir como prosa.\n\n"
        "2. CONSIDERACIONES Y ACCIONES DE MEJORA\n"
        "Presenta entre 2 y 3 áreas de crecimiento. Para cada una, integra en un mismo bloque: "
        "la consideración (redactada con verbos como Revisar, Ajustar, Fomentar, Fortalecer, "
        "Evaluar, Promover, Regular, Realizar), la evidencia que la sustenta (comentario o "
        "puntaje específico) y la acción SMART sugerida (específica, medible, alcanzable, "
        "relevante y temporal para la siguiente cohorte). Agrupa ideas similares, evita "
        "repeticiones y mantén el lenguaje en clave de \"áreas de crecimiento\", no de "
        "\"deficiencias\".\n\n"
        "3. RUTA DE FORMACIÓN PERSONALIZADA\n"
        "Con base en las áreas de crecimiento identificadas, diseña una ruta de formación de 2 a "
        "3 pasos, ordenados de mayor a menor prioridad. Para cada paso indica: el recurso de "
        "Aprende+ recomendado (nombre exacto, enlace y una frase que explique por qué es relevante "
        "para este docente en particular), y qué habilidad o resultado de aprendizaje específico "
        "del catálogo contribuye a resolver la necesidad detectada.\n\n"
        "Usa ÚNICAMENTE los recursos del siguiente catálogo institucional:\n\n"
        "TRAYECTORIAS:\n"
        "① Diseño de Experiencias de Aprendizaje\n"
        "   → Competencia: diseñar y liderar secuencias didácticas con el Ciclo de Kolb.\n"
        "   → Úsala cuando: el docente necesita estructurar mejor sus actividades, diversificar "
        "metodologías o promover la participación y autonomía del estudiante.\n\n"
        "② Trayectoria Innovación\n"
        "   → Competencia: aplicar herramientas de innovación, creatividad y metodologías ágiles.\n"
        "   → Úsala cuando: el docente busca renovar su práctica con enfoques más creativos o ágiles.\n\n"
        "③ Trayectoria Inteligencia Artificial\n"
        "   → Competencia: integrar IA para optimizar diseño de recursos, automatizar procesos y "
        "mejorar eficiencia.\n"
        "   → Úsala cuando: el docente necesita mejorar su gestión de recursos digitales o quiere "
        "innovar con IA en el aula.\n\n"
        "CURSOS INDIVIDUALES:\n"
        "④ Diseño de Syllabus y Rúbricas para la Evaluación del Aprendizaje\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/143311\n"
        "   → Úsalo cuando: hay debilidad en claridad de criterios de evaluación o estructura del curso.\n\n"
        "⑤ Metodología del Ciclo de Kolb\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/130033\n"
        "   → Úsalo cuando: el docente requiere una introducción al aprendizaje experiencial antes "
        "de abordar la trayectoria completa.\n\n"
        "⑥ Evaluación del Aprendizaje\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/184752\n"
        "   → Úsalo cuando: se detecta debilidad en estrategias de evaluación, retroalimentación "
        "o seguimiento del aprendizaje.\n\n"
        "⑦ Aprendizaje Basado en Retos (ABR)\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/134289\n"
        "   → Úsalo cuando: el docente quiere incorporar metodologías activas que promuevan "
        "pensamiento crítico y colaboración.\n\n"
        "⑧ Aprendizaje Basado en Proyectos (ABP)\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/127810\n"
        "   → Úsalo cuando: el docente quiere que los estudiantes aprendan resolviendo situaciones "
        "reales mediante proyectos colaborativos.\n\n"
        "⑨ El Pacto Pedagógico\n"
        "   Enlace: https://interactivavirtual.eafit.edu.co/d2l/le/discovery/view/course/202260\n"
        "   → Úsalo cuando: se detectan problemas de comunicación, acuerdos de convivencia o "
        "compromiso entre docente y estudiantes.\n\n"
        "No inventes recursos ni enlaces. Si ningún recurso del catálogo se ajusta a una necesidad "
        "detectada, indícalo explícitamente.\n\n"
        "PRINCIPIOS DE REDACCIÓN:\n"
        "- Objetivo formativo: énfasis en \"áreas de crecimiento\", no en \"deficiencias\".\n"
        "- Feedforward: orienta hacia acciones futuras, no hacia errores del pasado.\n"
        "- Equilibrio: lo positivo siempre antes que las oportunidades de mejora.\n"
        "- Tono profesional, respetuoso y de acompañamiento (nunca de juicio).\n"
        "- Coherencia, buena gramática y puntuación en todo el texto.\n"
        "- Sin redundancias ni extensiones innecesarias.\n"
        "- Usa SOLO los recursos del catálogo proporcionado; no inventes ni agregues otros."
    )

    # Cargar prompt guardado en GitHub (si existe) como valor inicial
    _prompt_guardado = None
    if _GH_TOKEN:
        try:
            _prompt_guardado = _gh_cargar_prompt_cached(_GH_TOKEN)
        except Exception:
            pass
    _prompt_inicial = _prompt_guardado if _prompt_guardado else PROMPT_BASE_USUARIO

    instruccion_usuaria = st.text_area(
        "Instrucciones para la IA",
        value=_prompt_inicial,
        height=380,
        help="Puedes modificar o agregar instrucciones. Este texto se envía directamente a la IA.",
        key="prompt_ia_editor",
    )

    col_guardar, col_restaurar = st.columns([2, 1])
    with col_guardar:
        if st.button("💾  Guardar prompt para siempre", use_container_width=True,
                     help="Guarda este prompt en GitHub. Se cargará automáticamente en todas las sesiones futuras."):
            if _GH_TOKEN:
                try:
                    with st.spinner("Guardando prompt…"):
                        ok = _gh_guardar_prompt(_GH_TOKEN, instruccion_usuaria)
                    if ok:
                        _gh_cargar_prompt_cached.clear()   # limpiar cache para que cargue el nuevo
                        st.success("✅ Prompt guardado. Se usará en todas las sesiones futuras.")
                    else:
                        st.error("❌ No se pudo guardar el prompt en GitHub.")
                except Exception as e:
                    st.error(f"❌ Error al guardar: {e}")
            else:
                st.warning("⚠️ Sin GITHUB_TOKEN configurado no se puede guardar.")
    with col_restaurar:
        if st.button("↩️  Restaurar por defecto", use_container_width=True,
                     help="Borra el prompt guardado y vuelve al prompt original del sistema."):
            if _GH_TOKEN:
                try:
                    sha = _gh_get_file_sha(_GH_TOKEN, _GH_PROMPT_FILE)
                    if sha:
                        url_del = f"https://api.github.com/repos/{_GH_REPO}/contents/{_GH_PROMPT_FILE}"
                        payload_del = _json.dumps({
                            "message": "Restaurar prompt por defecto",
                            "sha": sha, "branch": "main"
                        }).encode()
                        req_del = urllib.request.Request(url_del, data=payload_del,
                                                         method="DELETE", headers={
                            "Authorization": f"Bearer {_GH_TOKEN}",
                            "Accept": "application/vnd.github+json",
                            "Content-Type": "application/json",
                            "X-GitHub-Api-Version": "2022-11-28",
                        })
                        urllib.request.urlopen(req_del, timeout=15)
                    _gh_cargar_prompt_cached.clear()
                    st.success("✅ Prompt restaurado al valor por defecto. Recarga la página para verlo.")
                except Exception as e:
                    st.error(f"❌ Error al restaurar: {e}")
            else:
                st.warning("⚠️ Sin GITHUB_TOKEN configurado.")

    st.markdown('</div>', unsafe_allow_html=True)

    # ── Estado de sesión para el texto generado ──
    if "consideraciones_texto" not in st.session_state:
        st.session_state.consideraciones_texto = ""
    if "consideraciones_docx_bytes" not in st.session_state:
        st.session_state.consideraciones_docx_bytes = None
    if "consideraciones_nombre_archivo" not in st.session_state:
        st.session_state.consideraciones_nombre_archivo = None

    col_gen, col_regen = st.columns([1, 1])
    with col_gen:
        generar_clic = st.button("✨ Generar Consideraciones", use_container_width=True,
                                  disabled=(informe_bytes_base is None))
    with col_regen:
        regenerar_clic = st.button("🔁 Regenerar", use_container_width=True,
                                    disabled=(not st.session_state.consideraciones_texto))

    if informe_bytes_base is None:
        st.info("📄 Sube el informe (.docx) generado por la app para continuar.")

    if (generar_clic or regenerar_clic) and informe_bytes_base is not None:
        # Limpiar estado del informe anterior para evitar contaminación entre generaciones
        st.session_state.consideraciones_texto = ""
        st.session_state.consideraciones_docx_bytes = None
        st.session_state.consideraciones_nombre_archivo = None
        try:
            with st.spinner("Leyendo el informe y los documentos guía…"):
                informe_bytes = informe_bytes_base
                info_informe  = extraer_texto_informe_actual(informe_bytes)

                contexto_partes = []
                MAX_CHARS_POR_DOC = 4000   # ~1000 tokens por doc — gpt-4o tiene límite de 8k total

                # Documentos persistentes seleccionados (se descargan desde GitHub)
                for doc in docs_seleccionados_nombres:
                    try:
                        contenido = _gh_descargar_doc(doc["url"], _GH_TOKEN)
                        texto_doc = extraer_texto_referencia(doc["nombre"], contenido)
                        if texto_doc.strip():
                            texto_doc = texto_doc.strip()[:MAX_CHARS_POR_DOC]
                            contexto_partes.append(f"--- {doc['nombre']} ---\n{texto_doc}")
                    except Exception:
                        pass   # Si un doc no se puede descargar, se omite silenciosamente

                # Documentos nuevos subidos en esta sesión (ya en memoria)
                for nombre_doc, bytes_doc in docs_nuevos_bytes:
                    texto_doc = extraer_texto_referencia(nombre_doc, bytes_doc)
                    if texto_doc.strip():
                        texto_doc = texto_doc.strip()[:MAX_CHARS_POR_DOC]
                        contexto_partes.append(f"--- {nombre_doc} ---\n{texto_doc}")

                contexto_docs = "\n\n".join(contexto_partes)

            with st.spinner("Generando Consideraciones con IA…"):
                texto_generado = generar_consideraciones_ia(
                    _GH_MODELS_TOKEN, info_informe, contexto_docs, instruccion_usuaria,
                    nombre_docente=_nombre_docente_raw
                )

            st.session_state.consideraciones_texto = texto_generado
            st.session_state.consideraciones_docx_bytes = informe_bytes
            st.session_state.consideraciones_nombre_archivo = informe_nombre_base
            st.success("✅ Consideraciones generadas. Puedes revisarlas y editarlas abajo antes de descargar.")
        except RuntimeError as e:
            st.error(f"❌ {e}")
        except Exception as e:
            st.error(f"❌ Error inesperado al generar Consideraciones: {e}")

    # ── Previsualización y edición ──
    if st.session_state.consideraciones_texto:
        st.markdown('<div class="card"><div class="card-label">📝 Vista previa — editable</div>',
                    unsafe_allow_html=True)
        texto_editado = st.text_area(
            "Revisa, ajusta y valida el contenido antes de incorporarlo al informe",
            value=st.session_state.consideraciones_texto,
            height=260,
            key="texto_consideraciones_editor"
        )
        st.markdown('</div>', unsafe_allow_html=True)

        if st.button("✅ Incorporar al informe y preparar descarga", use_container_width=True):
            try:
                with st.spinner("Insertando Consideraciones en el informe…"):
                    docx_final = insertar_consideraciones_en_docx(
                        st.session_state.consideraciones_docx_bytes, texto_editado
                    )
                nombre_final = st.session_state.consideraciones_nombre_archivo
                st.session_state["_docx_final_bytes"] = docx_final
                st.session_state["_docx_final_nombre"] = nombre_final
            except RuntimeError as e:
                st.error(f"❌ {e}")
            except Exception as e:
                st.error(f"❌ Error al insertar Consideraciones en el informe: {e}")

    # Mostrar descarga si ya está lista
    if st.session_state.get("_docx_final_bytes"):
        st.success(f"✅ Informe listo: **{st.session_state['_docx_final_nombre']}**")
        col_dl, col_next = st.columns([2, 1])
        with col_dl:
            st.download_button(
                label="⬇️  Descargar informe con Consideraciones",
                data=st.session_state["_docx_final_bytes"],
                file_name=st.session_state["_docx_final_nombre"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with col_next:
            if st.button("➡️  Siguiente informe", use_container_width=True,
                         help="Limpia todo y deja la página lista para el siguiente docente."):
                for _k in ["_docx_final_bytes", "_docx_final_nombre",
                           "consideraciones_texto", "consideraciones_docx_bytes",
                           "consideraciones_nombre_archivo", "_nombre_docente_cache_key",
                           "_nombre_docente_cache", "_consideraciones_informe_key"]:
                    st.session_state.pop(_k, None)
                st.rerun()

